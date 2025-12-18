"""
Script to create comprehensive Word documentation for Multi-Tenant NLP2SQL System
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def add_heading_with_style(doc, text, level=1):
    """Add a styled heading"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_table_from_data(doc, headers, data):
    """Create a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True

    # Data rows
    for row_data in data:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = str(value)

    return table

def create_documentation():
    """Create the comprehensive documentation"""
    doc = Document()

    # ============================================================================
    # TITLE PAGE
    # ============================================================================
    title = doc.add_heading('Multi-Tenant NLP2SQL System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Comprehensive System Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(0, 0, 139)

    doc.add_paragraph()
    doc.add_paragraph()

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}\n')
    info_para.add_run('Version: 1.0\n')
    info_para.add_run('Location: D:\\Fundae\\Multi Tenant NLP2SQL')

    doc.add_page_break()

    # ============================================================================
    # TABLE OF CONTENTS (Manual)
    # ============================================================================
    add_heading_with_style(doc, 'Table of Contents', 1)

    toc_items = [
        '1. Executive Summary',
        '2. System Overview',
        '3. Architecture & Components',
        '4. Technology Stack',
        '5. Database Architecture',
        '6. Authentication & Security',
        '7. Multi-Tenancy Implementation',
        '8. NLP to SQL Engine',
        '9. API Endpoints',
        '10. Frontend Components',
        '11. Deployment & Configuration',
        '12. Testing Framework',
        '13. Data Models',
        '14. RBAC System',
        '15. Performance & Optimization',
        '16. Monitoring & Logging',
        '17. User Guide',
        '18. Appendices'
    ]

    for item in toc_items:
        doc.add_paragraph(item, style='List Number')

    doc.add_page_break()

    # ============================================================================
    # 1. EXECUTIVE SUMMARY
    # ============================================================================
    add_heading_with_style(doc, '1. Executive Summary', 1)

    doc.add_paragraph(
        'The Multi-Tenant NLP2SQL System is an enterprise-grade platform that enables users '
        'to query databases using natural language. The system provides complete tenant isolation, '
        'sophisticated role-based access control, and support for multiple database technologies '
        '(MySQL, PostgreSQL, MongoDB, and SQLite).'
    )

    add_heading_with_style(doc, 'Key Features', 2)
    features = [
        'Natural Language to SQL conversion using pattern matching and AI',
        'Complete multi-tenant isolation with separate databases per organization',
        'Advanced RBAC with hierarchical role templates',
        'Support for MySQL, PostgreSQL, MongoDB, and SQLite',
        'JWT-based authentication with session management',
        'Real-time query execution with visualization',
        'Automated tenant onboarding and database cloning',
        'Comprehensive security with SQL injection prevention',
        'Performance optimization with multi-level caching',
        'Human Digital Twin (HDT) personalization system',
        'Interactive Streamlit web interface',
        'RESTful FastAPI backend',
        'Docker containerization for easy deployment'
    ]

    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')

    add_heading_with_style(doc, 'System Metrics', 2)
    metrics_data = [
        ['Total Source Files', '35+ Python modules'],
        ['Lines of Code', '10,000+ lines'],
        ['Supported Databases', '4 types (MySQL, PostgreSQL, MongoDB, SQLite)'],
        ['Tenant Organizations', '5 pre-configured'],
        ['API Endpoints', '30+ REST endpoints'],
        ['Test Coverage', '53+ automated tests'],
        ['Role Templates', '6 hierarchical roles'],
        ['Industry Templates', '8+ pre-configured schemas']
    ]

    add_table_from_data(doc, ['Metric', 'Value'], metrics_data)

    doc.add_page_break()

    # ============================================================================
    # 2. SYSTEM OVERVIEW
    # ============================================================================
    add_heading_with_style(doc, '2. System Overview', 1)

    add_heading_with_style(doc, 'Purpose and Goals', 2)
    doc.add_paragraph(
        'The Multi-Tenant NLP2SQL System democratizes database access by allowing non-technical '
        'users to query complex databases using natural language. The system ensures complete '
        'data isolation between tenants while providing a unified administrative interface.'
    )

    add_heading_with_style(doc, 'Target Users', 2)
    users = [
        'Business Analysts - Query data without SQL knowledge',
        'Data Scientists - Rapid data exploration and analysis',
        'Executives - Access to real-time business intelligence',
        'System Administrators - Multi-tenant management',
        'Developers - API integration and customization'
    ]
    for user in users:
        doc.add_paragraph(user, style='List Bullet')

    add_heading_with_style(doc, 'Deployment Model', 2)
    doc.add_paragraph(
        'The system uses a microservices architecture deployed via Docker Compose, with separate '
        'containers for each tenant database, the main application, AI/NLP service (Ollama), and '
        'database administration tools.'
    )

    doc.add_page_break()

    # ============================================================================
    # 3. ARCHITECTURE & COMPONENTS
    # ============================================================================
    add_heading_with_style(doc, '3. Architecture & Components', 1)

    add_heading_with_style(doc, 'High-Level Architecture', 2)
    doc.add_paragraph(
        'The system follows a three-tier architecture with presentation (Streamlit), '
        'application (FastAPI), and data (Multi-database) layers.'
    )

    add_heading_with_style(doc, 'Directory Structure', 2)
    structure = [
        '/src - Core application logic (35 modules)',
        '/tests - Comprehensive testing suite (9 test files)',
        '/root_schemas - Database schema templates with versioning',
        '/database - Tenant database initialization scripts',
        '/rbac_schemas - Central RBAC database schema',
        '/databases - Runtime tenant SQLite databases',
        '/logs - Application logs and audit trails',
        '/exports - Query result exports (CSV, JSON, Excel)',
        '/demo_databases - Demo data stores'
    ]
    for item in structure:
        doc.add_paragraph(item, style='List Bullet')

    add_heading_with_style(doc, 'Core Components', 2)

    components_data = [
        ['Component', 'File', 'Purpose'],
        ['FastAPI Backend', 'src/main.py', 'REST API server on port 8000'],
        ['Streamlit Frontend', 'streamlit_app.py', 'Web UI on port 8501'],
        ['Auth Manager', 'src/auth.py', 'JWT authentication and password hashing'],
        ['RBAC Manager', 'src/tenant_rbac_manager.py', 'Role-based access control'],
        ['NLP2SQL Engine', 'src/nlp2sql_engine.py', 'Pattern-based NL to SQL'],
        ['Tenant NLP2SQL', 'src/tenant_aware_nlp2sql.py', 'Tenant-aware query engine'],
        ['Connection Manager', 'src/tenant_connection_manager.py', 'DB connection pooling'],
        ['Security Manager', 'src/security.py', 'SQL injection prevention'],
        ['Database Cloner', 'src/database_cloner.py', 'Tenant database provisioning'],
        ['HDT Manager', 'src/hdt_manager.py', 'Human Digital Twin system'],
        ['Performance Optimizer', 'src/performance_optimization.py', 'Caching & optimization'],
        ['Monitoring System', 'src/error_handling_monitoring.py', 'Error handling & alerts']
    ]

    add_table_from_data(doc, components_data[0], components_data[1:])

    doc.add_page_break()

    # ============================================================================
    # 4. TECHNOLOGY STACK
    # ============================================================================
    add_heading_with_style(doc, '4. Technology Stack', 1)

    add_heading_with_style(doc, 'Backend Technologies', 2)
    backend_tech = [
        ['Technology', 'Version', 'Purpose'],
        ['Python', '3.9+', 'Primary programming language'],
        ['FastAPI', '0.104.1', 'Modern async web framework'],
        ['Uvicorn', '0.24.0', 'ASGI server'],
        ['SQLAlchemy', '2.0.23', 'ORM and database abstraction'],
        ['Pydantic', '2.5.1', 'Data validation'],
        ['Python-dotenv', '1.0.0', 'Environment configuration']
    ]
    add_table_from_data(doc, backend_tech[0], backend_tech[1:])

    add_heading_with_style(doc, 'Frontend Technologies', 2)
    frontend_tech = [
        ['Technology', 'Version', 'Purpose'],
        ['Streamlit', '1.28.1', 'Interactive web interface'],
        ['Plotly', '5.17.0', 'Data visualization'],
        ['Pandas', '2.1.4', 'Data manipulation'],
        ['NumPy', '1.25.2', 'Numerical computing']
    ]
    add_table_from_data(doc, frontend_tech[0], frontend_tech[1:])

    add_heading_with_style(doc, 'Database Technologies', 2)
    db_tech = [
        ['Database', 'Version', 'Driver', 'Use Case'],
        ['MySQL', '8.0', 'mysql-connector-python 8.2.0', 'HealthPlus, EduLearn'],
        ['PostgreSQL', '15', 'psycopg2-binary 2.9.9', 'FinanceHub'],
        ['MongoDB', '7.0', 'pymongo 4.6.0', 'RetailMax'],
        ['SQLite', 'Built-in', 'sqlite3 (built-in)', 'TechCorp']
    ]
    add_table_from_data(doc, db_tech[0], db_tech[1:])

    add_heading_with_style(doc, 'Authentication & Security', 2)
    auth_tech = [
        ['Technology', 'Version', 'Purpose'],
        ['PyJWT', '2.8.0', 'JWT token encoding/decoding'],
        ['BCrypt', '4.1.2', 'Password hashing'],
        ['python-jose', '3.3.0', 'Cryptography operations'],
        ['Passlib', '1.7.4', 'Password utilities']
    ]
    add_table_from_data(doc, auth_tech[0], auth_tech[1:])

    add_heading_with_style(doc, 'NLP & AI', 2)
    nlp_tech = [
        ['Technology', 'Version', 'Purpose'],
        ['Ollama', '0.2.1', 'Local LLM integration (Llama 3.1)'],
        ['sqlparse', '0.4.4', 'SQL parsing and formatting']
    ]
    add_table_from_data(doc, nlp_tech[0], nlp_tech[1:])

    add_heading_with_style(doc, 'DevOps & Deployment', 2)
    devops_tech = [
        ['Technology', 'Purpose'],
        ['Docker', 'Containerization'],
        ['Docker Compose', 'Multi-container orchestration'],
        ['Pytest', 'Testing framework'],
        ['pytest-asyncio', 'Async test support']
    ]
    add_table_from_data(doc, devops_tech[0], devops_tech[1:])

    doc.add_page_break()

    # ============================================================================
    # 5. DATABASE ARCHITECTURE
    # ============================================================================
    add_heading_with_style(doc, '5. Database Architecture', 1)

    add_heading_with_style(doc, 'Multi-Database Strategy', 2)
    doc.add_paragraph(
        'The system implements true multi-tenancy with complete database isolation. '
        'Each tenant organization has its own dedicated database instance, ensuring '
        'zero data leakage and independent scaling.'
    )

    add_heading_with_style(doc, 'Tenant Databases', 2)
    tenant_dbs = [
        ['Tenant', 'Database Type', 'Port', 'Industry', 'Container'],
        ['TechCorp', 'SQLite', 'File-based', 'Technology', 'N/A (local file)'],
        ['HealthPlus', 'MySQL', '3307', 'Healthcare', 'healthplus_mysql'],
        ['FinanceHub', 'PostgreSQL', '5433', 'Finance', 'financehub_postgres'],
        ['RetailMax', 'MongoDB', '27018', 'Retail', 'retailmax_mongo'],
        ['EduLearn', 'MySQL', '3308', 'Education', 'edulearn_mysql']
    ]
    add_table_from_data(doc, tenant_dbs[0], tenant_dbs[1:])

    add_heading_with_style(doc, 'Central RBAC Database', 2)
    doc.add_paragraph(
        'A separate MySQL database (rbac_central) manages authentication, authorization, '
        'and cross-tenant user access. This database is NOT cloned and serves as the '
        'central control point for the entire system.'
    )

    rbac_tables = [
        ['Table Name', 'Purpose'],
        ['rbac_schema_info', 'Schema version tracking'],
        ['master_organizations', 'Tenant registry and metadata'],
        ['role_templates', 'System-wide role definitions'],
        ['master_users', 'Global user registry'],
        ['user_tenant_mappings', 'Cross-tenant access control'],
        ['user_tenant_roles', 'Role assignments per tenant'],
        ['user_permissions', 'Fine-grained permission matrix'],
        ['tenant_access_sessions', 'JWT session management'],
        ['rbac_audit_log', 'Complete audit trail']
    ]
    add_table_from_data(doc, rbac_tables[0], rbac_tables[1:])

    add_heading_with_style(doc, 'Tenant Database Schema (Common Structure)', 2)
    doc.add_paragraph(
        'All tenant databases follow a consistent structure with domain-specific tables. '
        'Key elements include:'
    )

    schema_elements = [
        'org_id column in every table for tenant isolation',
        'Foreign key relationships for data integrity',
        'Timestamps (created_at, updated_at) for auditing',
        'Appropriate indexes for performance',
        'Domain-specific entities (patients, accounts, products, etc.)'
    ]
    for element in schema_elements:
        doc.add_paragraph(element, style='List Bullet')

    add_heading_with_style(doc, 'HealthPlus Schema Example', 2)
    healthplus_schema = [
        ['Table', 'Primary Purpose', 'Key Columns'],
        ['patients', 'Patient master data', 'patient_id, first_name, last_name, date_of_birth'],
        ['doctors', 'Doctor registry', 'doctor_id, first_name, last_name, specialization'],
        ['treatments', 'Treatment catalog', 'treatment_id, treatment_name, base_cost'],
        ['appointments', 'Appointment scheduling', 'appointment_id, patient_id, doctor_id, appointment_date'],
        ['medical_records', 'Patient medical history', 'record_id, patient_id, diagnosis, prescription']
    ]
    add_table_from_data(doc, healthplus_schema[0], healthplus_schema[1:])

    add_heading_with_style(doc, 'Database Cloning System', 2)
    doc.add_paragraph(
        'The system includes an automated database cloning capability for rapid tenant '
        'onboarding. Root schema templates are maintained for each database type with '
        'versioning support.'
    )

    clone_features = [
        'Root schema templates in /root_schemas for each DB type',
        'Version tracking with migration scripts',
        'Automated cloning API (src/database_cloner.py)',
        'Clone verification and validation',
        'Rollback capabilities for failed clones',
        'Schema upgrade paths (v1.0 to v1.1)'
    ]
    for feature in clone_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_page_break()

    # ============================================================================
    # 6. AUTHENTICATION & SECURITY
    # ============================================================================
    add_heading_with_style(doc, '6. Authentication & Security', 1)

    add_heading_with_style(doc, 'Authentication Flow', 2)
    auth_steps = [
        '1. User submits email and password via Streamlit interface',
        '2. Email domain is used to identify organization (e.g., @healthplus.com → HealthPlus)',
        '3. Password is hashed with BCrypt and compared with stored hash',
        '4. JWT token is generated with user_id, org_id, and role',
        '5. Token is returned to client and stored in session',
        '6. All subsequent requests include JWT in Authorization header',
        '7. JWT middleware validates token on every request',
        '8. Token expiration is checked (default: 30 minutes)',
        '9. User can refresh token before expiration'
    ]
    for step in auth_steps:
        doc.add_paragraph(step, style='List Number')

    add_heading_with_style(doc, 'JWT Token Structure', 2)
    doc.add_paragraph('JWT tokens contain the following claims:')
    jwt_claims = [
        'sub: User ID (UUID)',
        'org_id: Organization ID',
        'email: User email address',
        'role: User role (admin, manager, analyst, developer, viewer)',
        'exp: Expiration timestamp',
        'iat: Issued at timestamp'
    ]
    for claim in jwt_claims:
        doc.add_paragraph(claim, style='List Bullet')

    add_heading_with_style(doc, 'Security Features', 2)

    security_features = [
        ['Feature', 'Implementation', 'Location'],
        ['Password Hashing', 'BCrypt with salt rounds', 'src/auth.py'],
        ['JWT Authentication', 'HS256 algorithm with secret key', 'src/jwt_middleware.py'],
        ['SQL Injection Prevention', 'Pattern matching and parameterized queries', 'src/security.py'],
        ['CORS Protection', 'FastAPI CORS middleware', 'src/main.py'],
        ['Rate Limiting', 'Request throttling per user', 'Configuration'],
        ['Session Management', 'Token expiration and refresh', 'src/tenant_rbac_manager.py'],
        ['Query Complexity Limits', 'Role-based query restrictions', 'Environment config'],
        ['Audit Logging', 'Complete query and access logs', 'rbac_audit_log table'],
        ['Data Isolation', 'org_id filtering on all queries', 'All database operations']
    ]
    add_table_from_data(doc, security_features[0], security_features[1:])

    add_heading_with_style(doc, 'SQL Injection Prevention', 2)
    doc.add_paragraph('The security manager implements multiple layers of protection:')

    sql_protection = [
        'Pattern-based detection of malicious SQL keywords',
        'Parameterized query enforcement',
        'Query type validation (SELECT, INSERT, UPDATE, DELETE)',
        'Comment stripping (-- and /* */)',
        'Semicolon terminator detection',
        'UNION, DROP, ALTER statement blocking',
        'System table access prevention'
    ]
    for item in sql_protection:
        doc.add_paragraph(item, style='List Bullet')

    add_heading_with_style(doc, 'Environment Variables Security', 2)
    doc.add_paragraph(
        'Sensitive configuration is stored in .env file (not committed to version control). '
        'Key security settings include:'
    )

    env_security = [
        'SECRET_KEY: JWT signing key (must be changed in production)',
        'Database passwords for all tenants',
        'RBAC database credentials',
        'Ollama API endpoint',
        'Rate limiting thresholds',
        'Query complexity limits per role'
    ]
    for item in env_security:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ============================================================================
    # 7. MULTI-TENANCY IMPLEMENTATION
    # ============================================================================
    add_heading_with_style(doc, '7. Multi-Tenancy Implementation', 1)

    add_heading_with_style(doc, 'Isolation Layers', 2)
    doc.add_paragraph('The system implements five layers of tenant isolation:')

    isolation_layers = [
        ['Layer', 'Mechanism', 'Implementation'],
        ['1. Physical Database', 'Separate database instance per tenant', 'Docker containers'],
        ['2. JWT Token', 'org_id embedded in token', 'jwt_middleware.py'],
        ['3. Routing Middleware', 'Request routing to correct DB', 'tenant_routing_middleware.py'],
        ['4. Connection Pool', 'Isolated connection pools', 'tenant_connection_manager.py'],
        ['5. Query Filter', 'WHERE org_id = ? on all queries', 'tenant_aware_nlp2sql.py']
    ]
    add_table_from_data(doc, isolation_layers[0], isolation_layers[1:])

    add_heading_with_style(doc, 'Request Lifecycle', 2)
    doc.add_paragraph('A typical query request flows through these stages:')

    request_flow = [
        '1. User submits natural language query via Streamlit UI',
        '2. Frontend sends HTTP POST to FastAPI with JWT token',
        '3. CORS middleware allows cross-origin request',
        '4. JWT middleware validates token and extracts tenant context',
        '5. Tenant routing middleware identifies target database',
        '6. RBAC manager checks user permissions for operation',
        '7. NLP2SQL engine converts query to SQL with tenant awareness',
        '8. Security manager validates SQL for injection attempts',
        '9. Query analyzer checks complexity against role limits',
        '10. Connection manager obtains connection from tenant pool',
        '11. SQL executed with org_id filter automatically applied',
        '12. Results collected and filtered by tenant',
        '13. Performance metrics recorded',
        '14. Query logged to audit trail',
        '15. Results formatted and returned to frontend',
        '16. Frontend visualizes data with Plotly charts'
    ]
    for step in request_flow:
        doc.add_paragraph(step, style='List Number')

    add_heading_with_style(doc, 'Tenant Identification', 2)
    doc.add_paragraph('Tenants are identified using multiple mechanisms:')

    tenant_id_methods = [
        'Email domain mapping (user@healthplus.com → HealthPlus)',
        'JWT token org_id claim',
        'Organization code in RBAC database',
        'Database name mapping in configuration'
    ]
    for method in tenant_id_methods:
        doc.add_paragraph(method, style='List Bullet')

    add_heading_with_style(doc, 'Cross-Tenant Access', 2)
    doc.add_paragraph(
        'The system supports users with access to multiple tenants through the '
        'CrossTenantUserManager component. Users can switch between authorized '
        'tenants without re-authentication.'
    )

    doc.add_page_break()

    # ============================================================================
    # 8. NLP TO SQL ENGINE
    # ============================================================================
    add_heading_with_style(doc, '8. NLP to SQL Engine', 1)

    add_heading_with_style(doc, 'Pattern-Based Approach', 2)
    doc.add_paragraph(
        'The NLP2SQL engine uses a sophisticated pattern-matching system to convert '
        'natural language queries into SQL. The system recognizes dozens of query patterns '
        'across multiple categories.'
    )

    add_heading_with_style(doc, 'Query Pattern Categories', 2)
    pattern_categories = [
        ['Category', 'Example Patterns', 'Generated SQL Type'],
        ['Basic Retrieval', '"show all customers", "list products"', 'SELECT * FROM table'],
        ['Count Queries', '"how many patients", "count doctors"', 'SELECT COUNT(*) FROM table'],
        ['Filtering', '"customers in New York", "products under $100"', 'SELECT with WHERE'],
        ['Aggregation', '"total sales by region", "average price"', 'GROUP BY with aggregates'],
        ['Sorting', '"top 10 customers by revenue"', 'ORDER BY with LIMIT'],
        ['Joins', '"patients with their doctors"', 'INNER/LEFT JOIN'],
        ['Date Ranges', '"appointments this month", "sales last year"', 'Date-based WHERE'],
        ['Business Intelligence', '"best selling products", "high value customers"', 'Complex aggregations']
    ]
    add_table_from_data(doc, pattern_categories[0], pattern_categories[1:])

    add_heading_with_style(doc, 'Tenant-Aware NLP2SQL', 2)
    doc.add_paragraph('The TenantAwareNLP2SQL component provides:')

    tenant_nlp_features = [
        'Schema caching per tenant for performance',
        'Automatic org_id filtering in generated SQL',
        'Table and column name validation against tenant schema',
        'Query type classification (SELECT, JOIN, AGGREGATE)',
        'Security level determination (SAFE, RESTRICTED, ADMIN_ONLY)',
        'Query complexity scoring',
        'Execution time tracking',
        'Result metadata capture'
    ]
    for feature in tenant_nlp_features:
        doc.add_paragraph(feature, style='List Bullet')

    add_heading_with_style(doc, 'Query Processing Pipeline', 2)
    query_pipeline = [
        '1. Input Analysis - Identify query pattern and extract entities',
        '2. Schema Validation - Verify table and column names exist',
        '3. SQL Generation - Apply pattern template with parameters',
        '4. Security Check - Validate for SQL injection attempts',
        '5. Permission Check - Verify user role allows query type',
        '6. Complexity Check - Ensure within role limits',
        '7. Tenant Filter - Add WHERE org_id = ? automatically',
        '8. Optimization - Apply caching and index hints',
        '9. Execution - Run against tenant database',
        '10. Result Processing - Format and return data'
    ]
    for step in query_pipeline:
        doc.add_paragraph(step, style='List Number')

    add_heading_with_style(doc, 'Example Query Transformations', 2)

    query_examples = [
        ['Natural Language', 'Generated SQL (simplified)'],
        ['"show all patients"', 'SELECT * FROM patients WHERE org_id = ?'],
        ['"how many doctors in cardiology"', 'SELECT COUNT(*) FROM doctors WHERE specialization = ? AND org_id = ?'],
        ['"top 5 customers by sales"', 'SELECT * FROM customers ORDER BY total_sales DESC LIMIT 5 WHERE org_id = ?'],
        ['"average treatment cost"', 'SELECT AVG(base_cost) FROM treatments WHERE org_id = ?'],
        ['"appointments this month"', 'SELECT * FROM appointments WHERE MONTH(appointment_date) = ? AND org_id = ?']
    ]
    add_table_from_data(doc, query_examples[0], query_examples[1:])

    doc.add_page_break()

    # ============================================================================
    # 9. API ENDPOINTS
    # ============================================================================
    add_heading_with_style(doc, '9. API Endpoints', 1)

    add_heading_with_style(doc, 'Authentication Endpoints', 2)
    auth_endpoints = [
        ['Method', 'Endpoint', 'Description', 'Auth Required'],
        ['POST', '/auth/login', 'User login with email/password', 'No'],
        ['POST', '/auth/register', 'New user registration', 'No'],
        ['GET', '/auth/verify', 'Verify JWT token validity', 'Yes'],
        ['POST', '/auth/refresh', 'Refresh expired token', 'Yes'],
        ['POST', '/auth/logout', 'Logout and invalidate token', 'Yes']
    ]
    add_table_from_data(doc, auth_endpoints[0], auth_endpoints[1:])

    add_heading_with_style(doc, 'Query Execution Endpoints', 2)
    query_endpoints = [
        ['Method', 'Endpoint', 'Description', 'Auth Required'],
        ['POST', '/api/query', 'Execute natural language query', 'Yes'],
        ['GET', '/api/query/suggestions', 'Get query suggestions', 'Yes'],
        ['GET', '/api/query/history', 'Get user query history', 'Yes'],
        ['POST', '/api/query/export', 'Export query results (CSV/JSON/Excel)', 'Yes']
    ]
    add_table_from_data(doc, query_endpoints[0], query_endpoints[1:])

    add_heading_with_style(doc, 'Multi-Tenant Management', 2)
    tenant_endpoints = [
        ['Method', 'Endpoint', 'Description', 'Auth Required'],
        ['POST', '/api/v1/onboarding/register-tenant', 'Register new tenant organization', 'Admin'],
        ['GET', '/api/v1/onboarding/status/{workflow_id}', 'Check onboarding workflow status', 'Admin'],
        ['GET', '/api/v1/tenants', 'List all tenant organizations', 'Admin'],
        ['GET', '/api/v1/tenants/{tenant_id}', 'Get tenant details', 'Admin']
    ]
    add_table_from_data(doc, tenant_endpoints[0], tenant_endpoints[1:])

    add_heading_with_style(doc, 'RBAC Endpoints', 2)
    rbac_endpoints = [
        ['Method', 'Endpoint', 'Description', 'Auth Required'],
        ['POST', '/api/v1/rbac/users', 'Create new user', 'Admin'],
        ['GET', '/api/v1/rbac/users/{user_id}', 'Get user profile and permissions', 'Yes'],
        ['POST', '/api/v1/rbac/roles', 'Assign role to user', 'Admin'],
        ['GET', '/api/v1/rbac/permissions', 'Check user permissions', 'Yes'],
        ['POST', '/api/v1/rbac/sessions', 'Create new session', 'Yes'],
        ['GET', '/api/v1/rbac/audit-log', 'Get audit log entries', 'Admin']
    ]
    add_table_from_data(doc, rbac_endpoints[0], rbac_endpoints[1:])

    add_heading_with_style(doc, 'Database Cloning API', 2)
    clone_endpoints = [
        ['Method', 'Endpoint', 'Description', 'Auth Required'],
        ['POST', '/api/v1/cloning/clone-from-root', 'Clone new tenant database', 'Admin'],
        ['GET', '/api/v1/cloning/clone-status/{clone_id}', 'Get clone operation progress', 'Admin'],
        ['GET', '/api/v1/cloning/system/status', 'Get cloning system status', 'Admin']
    ]
    add_table_from_data(doc, clone_endpoints[0], clone_endpoints[1:])

    add_heading_with_style(doc, 'Schema & Connection Endpoints', 2)
    schema_endpoints = [
        ['Method', 'Endpoint', 'Description', 'Auth Required'],
        ['GET', '/api/v1/schema/{tenant_id}', 'Get tenant database schema', 'Yes'],
        ['GET', '/api/v1/connections/{tenant_id}', 'Get connection pool health', 'Admin'],
        ['GET', '/api/v1/connections/{tenant_id}/metrics', 'Get connection pool metrics', 'Admin']
    ]
    add_table_from_data(doc, schema_endpoints[0], schema_endpoints[1:])

    doc.add_page_break()

    # ============================================================================
    # 10. FRONTEND COMPONENTS
    # ============================================================================
    add_heading_with_style(doc, '10. Frontend Components', 1)

    add_heading_with_style(doc, 'Streamlit Web Interface', 2)
    doc.add_paragraph(
        'The primary user interface is built with Streamlit, providing an interactive '
        'web application accessible via browser on port 8501.'
    )

    add_heading_with_style(doc, 'Main Interface Sections', 2)

    ui_sections = [
        ['Section', 'Functionality', 'User Interaction'],
        ['Login Panel', 'Email/password authentication', 'Form with submit button'],
        ['Sidebar', 'User info, logout, navigation', 'Always visible when logged in'],
        ['Query Input', 'Natural language query entry', 'Text area with submit'],
        ['Results Display', 'Tabular data presentation', 'Paginated table view'],
        ['Visualization', 'Charts and graphs', 'Interactive Plotly visualizations'],
        ['Export Options', 'Download results', 'CSV, JSON, Excel buttons'],
        ['Query History', 'Past queries and results', 'Scrollable list with re-run'],
        ['Dashboard', 'User analytics and stats', 'Metrics and summaries']
    ]
    add_table_from_data(doc, ui_sections[0], ui_sections[1:])

    add_heading_with_style(doc, 'Visualization Capabilities', 2)
    doc.add_paragraph('The system automatically generates appropriate visualizations:')

    viz_types = [
        'Bar Charts - For categorical comparisons',
        'Line Charts - For time-series data',
        'Scatter Plots - For correlation analysis',
        'Pie Charts - For distribution analysis',
        'Data Tables - For detailed row-level inspection',
        'Summary Statistics - For numerical data overview'
    ]
    for viz in viz_types:
        doc.add_paragraph(viz, style='List Bullet')

    add_heading_with_style(doc, 'Session Management', 2)
    doc.add_paragraph('Streamlit session state manages:')

    session_items = [
        'authenticated: Boolean flag for login status',
        'access_token: JWT token for API requests',
        'user_info: User profile and organization data',
        'chat_history: Query and response history',
        'current_results: Latest query results',
        'selected_visualization: Current chart type'
    ]
    for item in session_items:
        doc.add_paragraph(item, style='List Bullet')

    add_heading_with_style(doc, 'Additional Frontend Files', 2)

    frontend_files = [
        ['File', 'Purpose'],
        ['streamlit_app.py', 'Main production interface'],
        ['frontend_professional.py', 'Professional styled interface with advanced features'],
        ['working_frontend.py', 'Alternative simplified interface'],
        ['run_streamlit.py', 'Streamlit launcher script'],
        ['run_dashboard.py', 'Admin dashboard launcher'],
        ['security_validation_dashboard.py', 'Security monitoring interface']
    ]
    add_table_from_data(doc, frontend_files[0], frontend_files[1:])

    doc.add_page_break()

    # ============================================================================
    # 11. DEPLOYMENT & CONFIGURATION
    # ============================================================================
    add_heading_with_style(doc, '11. Deployment & Configuration', 1)

    add_heading_with_style(doc, 'Docker Compose Architecture', 2)
    doc.add_paragraph('The system deploys via Docker Compose with the following services:')

    docker_services = [
        ['Service', 'Container Name', 'Port', 'Purpose'],
        ['healthplus-mysql', 'healthplus_mysql', '3307→3306', 'MySQL for HealthPlus tenant'],
        ['financehub-postgres', 'financehub_postgres', '5433→5432', 'PostgreSQL for FinanceHub'],
        ['retailmax-mongo', 'retailmax_mongo', '27018→27017', 'MongoDB for RetailMax'],
        ['edulearn-mysql', 'edulearn_mysql', '3308→3306', 'MySQL for EduLearn'],
        ['nlp2sql-app', 'nlp2sql_app', '8001, 8501', 'Main application (FastAPI + Streamlit)'],
        ['ollama', 'ollama_nlp', '11434', 'AI/NLP service (Llama 3.1)'],
        ['phpmyadmin', 'phpmyadmin_healthplus', '8080', 'MySQL admin interface'],
        ['pgadmin', 'pgadmin_financehub', '8081', 'PostgreSQL admin interface'],
        ['mongo-express', 'mongo_express_retailmax', '8082', 'MongoDB admin interface']
    ]
    add_table_from_data(doc, docker_services[0], docker_services[1:])

    add_heading_with_style(doc, 'Environment Configuration', 2)
    doc.add_paragraph('Key configuration variables in .env file:')

    env_config = [
        ['Category', 'Key Variables', 'Purpose'],
        ['Security', 'SECRET_KEY, ALGORITHM, ACCESS_TOKEN_EXPIRE_MINUTES', 'JWT configuration'],
        ['Metadata DB', 'MYSQL_HOST, MYSQL_PORT, MYSQL_USER, MYSQL_PASSWORD', 'Central metadata database'],
        ['TechCorp', 'TECHCORP_DB_TYPE, TECHCORP_DB_NAME', 'SQLite configuration'],
        ['HealthPlus', 'HEALTHPLUS_DB_HOST, HEALTHPLUS_DB_PORT, HEALTHPLUS_DB_USER', 'MySQL connection'],
        ['FinanceHub', 'FINANCEHUB_DB_HOST, FINANCEHUB_DB_PORT, FINANCEHUB_DB_USER', 'PostgreSQL connection'],
        ['RetailMax', 'RETAILMAX_DB_HOST, RETAILMAX_DB_PORT, RETAILMAX_DB_USER', 'MongoDB connection'],
        ['EduLearn', 'EDULEARN_DB_HOST, EDULEARN_DB_PORT, EDULEARN_DB_USER', 'MySQL connection'],
        ['AI/NLP', 'OLLAMA_BASE_URL, OLLAMA_MODEL', 'LLM service'],
        ['Application', 'FASTAPI_HOST, FASTAPI_PORT, STREAMLIT_PORT, DEBUG', 'Server settings'],
        ['Security Limits', 'RATE_LIMIT_REQUESTS, MAX_QUERY_COMPLEXITY_*', 'Rate and complexity limits']
    ]
    add_table_from_data(doc, env_config[0], env_config[1:])

    add_heading_with_style(doc, 'Deployment Steps', 2)

    deployment_steps = [
        '1. Clone repository to deployment server',
        '2. Configure .env file with production credentials',
        '3. Update SECRET_KEY to cryptographically secure random value',
        '4. Start Docker Compose: docker-compose up -d',
        '5. Wait for all health checks to pass',
        '6. Initialize RBAC database with role templates',
        '7. Create initial admin user',
        '8. Pull Ollama model: docker exec ollama_nlp ollama pull llama3.1',
        '9. Verify all services are running: docker-compose ps',
        '10. Access application at http://localhost:8501'
    ]
    for step in deployment_steps:
        doc.add_paragraph(step, style='List Number')

    add_heading_with_style(doc, 'Production Considerations', 2)

    production_items = [
        'Use strong SECRET_KEY (256-bit random value)',
        'Enable HTTPS with SSL certificates',
        'Set DEBUG=False in production',
        'Configure database backups for all tenant databases',
        'Set up log rotation for /logs directory',
        'Implement monitoring and alerting',
        'Configure firewall rules to restrict access',
        'Use secrets management (e.g., Docker secrets, AWS Secrets Manager)',
        'Set appropriate resource limits in docker-compose.yml',
        'Enable database replication for high availability'
    ]
    for item in production_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ============================================================================
    # 12. TESTING FRAMEWORK
    # ============================================================================
    add_heading_with_style(doc, '12. Testing Framework', 1)

    add_heading_with_style(doc, 'Test Organization', 2)
    doc.add_paragraph('The /tests directory contains comprehensive test suites:')

    test_files = [
        ['Test File', 'Purpose', 'Coverage'],
        ['conftest.py', 'Pytest fixtures and configuration', 'Test infrastructure'],
        ['test_comprehensive_suite.py', 'End-to-end integration tests', 'Full system workflows'],
        ['test_security.py', 'Security validation tests', 'Auth, RBAC, SQL injection'],
        ['tenant_isolation_tester.py', 'Cross-tenant isolation verification', 'Data leakage prevention'],
        ['nlp2sql_accuracy_tester.py', 'NLP conversion accuracy', 'Query pattern matching'],
        ['security_penetration_tester.py', 'Attack simulation', 'SQL injection, XSS'],
        ['load_testing_framework.py', 'Performance under load', 'Concurrent users'],
        ['integration_testing_suite.py', 'Component integration', 'API, DB, NLP'],
        ['test_mock_comprehensive.py', 'Mock-based unit tests', 'Individual components']
    ]
    add_table_from_data(doc, test_files[0], test_files[1:])

    add_heading_with_style(doc, 'Test Categories (Pytest Markers)', 2)

    test_markers = [
        ['Marker', 'Command', 'Purpose'],
        ['security', 'pytest -m security', 'Run security tests only'],
        ['load', 'pytest -m load', 'Run load/performance tests'],
        ['isolation', 'pytest -m isolation', 'Run tenant isolation tests'],
        ['nlp2sql', 'pytest -m nlp2sql', 'Run NLP conversion tests'],
        ['e2e', 'pytest -m e2e', 'Run end-to-end tests']
    ]
    add_table_from_data(doc, test_markers[0], test_markers[1:])

    add_heading_with_style(doc, 'Test Coverage', 2)
    doc.add_paragraph('Pytest is configured with coverage reporting:')

    coverage_items = [
        'HTML coverage reports in htmlcov/',
        'XML coverage for CI/CD integration',
        'Terminal coverage summary',
        'Target: >80% code coverage',
        'Excludes: Test files, migrations, demo scripts'
    ]
    for item in coverage_items:
        doc.add_paragraph(item, style='List Bullet')

    add_heading_with_style(doc, 'Running Tests', 2)

    test_commands = [
        ['Command', 'Description'],
        ['pytest', 'Run all tests'],
        ['pytest tests/test_security.py', 'Run specific test file'],
        ['pytest -m security', 'Run tests with marker'],
        ['pytest --cov=src --cov-report=html', 'Run with coverage report'],
        ['pytest -v', 'Verbose output'],
        ['pytest -x', 'Stop on first failure'],
        ['pytest --lf', 'Run last failed tests only']
    ]
    add_table_from_data(doc, test_commands[0], test_commands[1:])

    doc.add_page_break()

    # ============================================================================
    # 13. DATA MODELS
    # ============================================================================
    add_heading_with_style(doc, '13. Data Models', 1)

    add_heading_with_style(doc, 'SQLAlchemy ORM Models', 2)
    doc.add_paragraph('The src/models.py file defines core data models:')

    models = [
        ['Model Class', 'Table', 'Purpose', 'Key Fields'],
        ['Organization', 'organizations', 'Tenant organizations', 'org_id, org_name, domain, database_type'],
        ['User', 'users', 'User accounts', 'user_id, org_id, email, role, password_hash'],
        ['HumanDigitalTwin', 'human_digital_twins', 'HDT profiles', 'hdt_id, name, skillset, languages'],
        ['Agent', 'agents', 'AI agents', 'agent_id, agent_type, capabilities, config'],
        ['HDTAgent', 'hdt_agents', 'HDT-Agent mapping', 'hdt_id, agent_id'],
        ['UserHDTAssignment', 'user_hdt_assignments', 'User-HDT links', 'user_id, hdt_id'],
        ['UserPermission', 'user_permissions', 'Fine-grained permissions', 'user_id, resource_type, access_level'],
        ['QueryLog', 'query_log', 'Query audit trail', 'query_text, generated_sql, execution_status'],
        ['UserSession', 'user_sessions', 'Active sessions', 'session_id, user_id, token_hash, expires_at']
    ]
    add_table_from_data(doc, models[0], models[1:])

    add_heading_with_style(doc, 'Pydantic Models (API)', 2)
    doc.add_paragraph('Request/response models for API validation:')

    pydantic_models = [
        ['Model', 'Purpose', 'Fields'],
        ['LoginRequest', 'Login credentials', 'email, password'],
        ['LoginResponse', 'Login result', 'access_token, user, organization, hdt'],
        ['QueryRequest', 'Query submission', 'query_text, export_format'],
        ['QueryResponse', 'Query results', 'query_id, generated_sql, results, execution_time'],
        ['UserResponse', 'User profile', 'user_id, username, email, role, org_id'],
        ['OrganizationResponse', 'Org details', 'org_id, org_name, domain, database_type'],
        ['HDTResponse', 'HDT profile', 'hdt_id, name, skillset, languages, agents'],
        ['AgentResponse', 'Agent details', 'agent_id, agent_name, agent_type, capabilities']
    ]
    add_table_from_data(doc, pydantic_models[0], pydantic_models[1:])

    add_heading_with_style(doc, 'Enums', 2)

    enums = [
        ['Enum', 'Values', 'Usage'],
        ['UserRole', 'admin, manager, analyst, developer, viewer', 'User role assignment'],
        ['AgentType', 'nlp2sql, rag, analytics, reporting, chatbot', 'AI agent classification'],
        ['DatabaseType', 'mysql, postgresql, mongodb, sqlite', 'Database technology'],
        ['AccessLevel', 'read, write, admin', 'Permission levels']
    ]
    add_table_from_data(doc, enums[0], enums[1:])

    doc.add_page_break()

    # ============================================================================
    # 14. RBAC SYSTEM
    # ============================================================================
    add_heading_with_style(doc, '14. Role-Based Access Control (RBAC)', 1)

    add_heading_with_style(doc, 'Role Hierarchy', 2)
    doc.add_paragraph('The system implements a hierarchical role structure:')

    role_hierarchy = [
        ['Role', 'Level', 'Key Permissions', 'Query Complexity Limit'],
        ['SUPER_ADMIN', '0 (Global)', 'System-wide administration, tenant management', '50'],
        ['Admin', '1 (Tenant)', 'Full tenant access, user management, all queries', '50'],
        ['Manager', '2', 'User oversight, reporting, complex queries', '20'],
        ['Analyst', '3', 'Data analysis, custom queries, exports', '15'],
        ['Developer', '4', 'Query creation, API access', '10'],
        ['Viewer', '5', 'Read-only, pre-defined queries', '5'],
        ['Guest', '6', 'Temporary limited access', '3']
    ]
    add_table_from_data(doc, role_hierarchy[0], role_hierarchy[1:])

    add_heading_with_style(doc, 'Role Templates', 2)
    doc.add_paragraph('Role templates define permissions as JSON structures:')

    role_template_structure = [
        'role_name: Unique identifier (e.g., "admin", "analyst")',
        'display_name: Human-readable name',
        'permissions: JSON array of permission objects',
        'inherits_from: Parent role for inheritance',
        'is_system_role: Whether role is built-in',
        'is_assignable: Whether role can be assigned to users',
        'metadata: Additional role configuration'
    ]
    for item in role_template_structure:
        doc.add_paragraph(item, style='List Bullet')

    add_heading_with_style(doc, 'Permission Types', 2)

    permission_types = [
        ['Permission', 'Scope', 'Roles'],
        ['query.execute', 'Execute SELECT queries', 'All roles'],
        ['query.create', 'Create custom queries', 'Analyst, Developer, Manager, Admin'],
        ['query.delete', 'Delete query history', 'Manager, Admin'],
        ['data.export', 'Export query results', 'Analyst, Developer, Manager, Admin'],
        ['user.create', 'Create new users', 'Manager, Admin'],
        ['user.delete', 'Delete users', 'Admin only'],
        ['tenant.manage', 'Tenant administration', 'SUPER_ADMIN only'],
        ['rbac.assign', 'Assign roles', 'Manager, Admin'],
        ['schema.view', 'View database schema', 'Developer, Manager, Admin'],
        ['audit.view', 'View audit logs', 'Admin, SUPER_ADMIN']
    ]
    add_table_from_data(doc, permission_types[0], permission_types[1:])

    add_heading_with_style(doc, 'RBAC Manager Components', 2)

    rbac_components = [
        ['Component', 'File', 'Responsibility'],
        ['TenantRBACManager', 'tenant_rbac_manager.py', 'Core RBAC engine, session management'],
        ['RoleTemplateManager', 'rbac_role_templates.py', 'Role template installation and management'],
        ['CrossTenantUserManager', 'cross_tenant_user_manager.py', 'Multi-tenant user access'],
        ['JWTMiddleware', 'jwt_middleware.py', 'Request authentication and authorization'],
        ['RBACDependencies', 'jwt_middleware.py', 'FastAPI dependency injection for auth']
    ]
    add_table_from_data(doc, rbac_components[0], rbac_components[1:])

    add_heading_with_style(doc, 'Permission Checking Flow', 2)

    permission_check = [
        '1. User makes authenticated request with JWT',
        '2. JWT middleware extracts user_id, org_id, and role',
        '3. RBACManager queries user_permissions table',
        '4. Permission inheritance applied from role hierarchy',
        '5. Resource-specific permissions checked',
        '6. Access level validated (read, write, admin)',
        '7. Result cached for performance',
        '8. Request allowed or denied with 403 Forbidden'
    ]
    for step in permission_check:
        doc.add_paragraph(step, style='List Number')

    doc.add_page_break()

    # ============================================================================
    # 15. PERFORMANCE & OPTIMIZATION
    # ============================================================================
    add_heading_with_style(doc, '15. Performance & Optimization', 1)

    add_heading_with_style(doc, 'Multi-Level Caching', 2)
    doc.add_paragraph('The PerformanceOptimizer implements three caching layers:')

    caching_layers = [
        ['Cache Type', 'Scope', 'TTL', 'Purpose'],
        ['Schema Cache', 'Per tenant', 'Long (hours)', 'Table and column metadata'],
        ['Query Result Cache', 'Per user', 'Medium (minutes)', 'Identical query results'],
        ['Permission Cache', 'Per user', 'Short (seconds)', 'RBAC permission checks'],
        ['Connection Pool Cache', 'Per tenant', 'Persistent', 'Database connections']
    ]
    add_table_from_data(doc, caching_layers[0], caching_layers[1:])

    add_heading_with_style(doc, 'Connection Pooling', 2)
    doc.add_paragraph('TenantConnectionManager provides:')

    pool_features = [
        'Separate connection pools per tenant database',
        'Configurable pool size (min/max connections)',
        'Connection health checking and recycling',
        'Automatic failover and retry logic',
        'Connection timeout handling',
        'Pool metrics and monitoring',
        'Graceful pool shutdown on tenant removal'
    ]
    for feature in pool_features:
        doc.add_paragraph(feature, style='List Bullet')

    add_heading_with_style(doc, 'Query Optimization', 2)

    query_optimization = [
        ['Technique', 'Implementation', 'Benefit'],
        ['Prepared Statements', 'Parameterized queries throughout', 'Performance + security'],
        ['Index Hints', 'Automatic index usage for common patterns', 'Faster lookups'],
        ['Result Limiting', 'MAX_QUERY_RESULTS configuration', 'Memory efficiency'],
        ['Lazy Loading', 'Pagination for large result sets', 'Reduced latency'],
        ['Query Complexity Analysis', 'Block expensive queries for low roles', 'Resource protection'],
        ['Connection Reuse', 'Pooling instead of per-query connections', 'Reduced overhead']
    ]
    add_table_from_data(doc, query_optimization[0], query_optimization[1:])

    add_heading_with_style(doc, 'Performance Metrics', 2)
    doc.add_paragraph('The system tracks:')

    metrics = [
        'Query execution time (milliseconds)',
        'Rows returned per query',
        'Cache hit/miss ratios',
        'Connection pool utilization',
        'API endpoint response times',
        'Database CPU and memory usage',
        'Concurrent user count',
        'Request rate per tenant'
    ]
    for metric in metrics:
        doc.add_paragraph(metric, style='List Bullet')

    doc.add_page_break()

    # ============================================================================
    # 16. MONITORING & LOGGING
    # ============================================================================
    add_heading_with_style(doc, '16. Monitoring & Logging', 1)

    add_heading_with_style(doc, 'Logging System', 2)
    doc.add_paragraph('Comprehensive logging is implemented throughout the system:')

    log_levels = [
        ['Level', 'Usage', 'Examples'],
        ['DEBUG', 'Development troubleshooting', 'SQL queries, cache operations'],
        ['INFO', 'Normal operations', 'User login, query execution, API calls'],
        ['WARNING', 'Potential issues', 'Slow queries, cache misses, deprecated features'],
        ['ERROR', 'Error conditions', 'Failed queries, authentication errors, DB connection issues'],
        ['CRITICAL', 'System failures', 'Database unavailable, service crashes']
    ]
    add_table_from_data(doc, log_levels[0], log_levels[1:])

    add_heading_with_style(doc, 'Log Files', 2)
    doc.add_paragraph('Logs are written to /logs directory:')

    log_files = [
        'app.log - Main application log',
        'query.log - All query executions',
        'auth.log - Authentication attempts',
        'security.log - Security events and violations',
        'performance.log - Performance metrics',
        'error.log - All errors and exceptions'
    ]
    for log_file in log_files:
        doc.add_paragraph(log_file, style='List Bullet')

    add_heading_with_style(doc, 'Audit Trail', 2)
    doc.add_paragraph('The rbac_audit_log table captures:')

    audit_items = [
        'All user logins and logouts',
        'Query executions (successful and failed)',
        'Role assignments and changes',
        'Permission grants and revocations',
        'Tenant creation and modification',
        'User creation and deletion',
        'Configuration changes',
        'Security violations and blocked queries'
    ]
    for item in audit_items:
        doc.add_paragraph(item, style='List Bullet')

    add_heading_with_style(doc, 'Monitoring System', 2)
    doc.add_paragraph('The error_handling_monitoring.py module provides:')

    monitoring_features = [
        ['Feature', 'Capability'],
        ['Health Checks', 'Periodic database connectivity verification'],
        ['Resource Monitoring', 'CPU, memory, disk usage tracking'],
        ['Alert System', 'Email/webhook notifications for critical events'],
        ['Uptime Tracking', 'Service availability metrics'],
        ['Error Rate Monitoring', 'Track error frequency per endpoint'],
        ['Slow Query Detection', 'Identify performance bottlenecks'],
        ['Connection Pool Metrics', 'Monitor pool health and saturation']
    ]
    add_table_from_data(doc, monitoring_features[0], monitoring_features[1:])

    add_heading_with_style(doc, 'Dashboard Access', 2)

    dashboards = [
        ['Dashboard', 'URL', 'Purpose'],
        ['Main Application', 'http://localhost:8501', 'User interface'],
        ['FastAPI Docs', 'http://localhost:8000/docs', 'API documentation'],
        ['phpMyAdmin', 'http://localhost:8080', 'MySQL administration'],
        ['pgAdmin', 'http://localhost:8081', 'PostgreSQL administration'],
        ['Mongo Express', 'http://localhost:8082', 'MongoDB administration'],
        ['Security Dashboard', 'security_validation_dashboard.py', 'Security monitoring'],
        ['Tenant Dashboard', 'src/tenant_management_dashboard.py', 'Tenant administration']
    ]
    add_table_from_data(doc, dashboards[0], dashboards[1:])

    doc.add_page_break()

    # ============================================================================
    # 17. USER GUIDE
    # ============================================================================
    add_heading_with_style(doc, '17. User Guide', 1)

    add_heading_with_style(doc, 'Getting Started', 2)

    getting_started = [
        '1. Ensure Docker and Docker Compose are installed',
        '2. Navigate to project directory: D:\\Fundae\\Multi Tenant NLP2SQL',
        '3. Start services: docker-compose up -d',
        '4. Wait for health checks to complete (~2 minutes)',
        '5. Open browser to http://localhost:8501',
        '6. Login with demo credentials (see Pre-configured Users below)',
        '7. Enter natural language query',
        '8. View results and visualizations'
    ]
    for step in getting_started:
        doc.add_paragraph(step, style='List Number')

    add_heading_with_style(doc, 'Pre-configured Demo Users', 2)

    demo_users = [
        ['Email', 'Password', 'Organization', 'Role', 'Database'],
        ['admin@techcorp.com', 'admin123', 'TechCorp', 'Admin', 'SQLite'],
        ['admin@healthplus.com', 'admin123', 'HealthPlus', 'Admin', 'MySQL (3307)'],
        ['doctor@healthplus.com', 'doctor123', 'HealthPlus', 'Analyst', 'MySQL (3307)'],
        ['admin@financehub.net', 'admin123', 'FinanceHub', 'Admin', 'PostgreSQL (5433)'],
        ['admin@retailmax.com', 'admin123', 'RetailMax', 'Admin', 'MongoDB (27018)'],
        ['admin@edulearn.org', 'admin123', 'EduLearn', 'Admin', 'MySQL (3308)']
    ]
    add_table_from_data(doc, demo_users[0], demo_users[1:])

    add_heading_with_style(doc, 'Example Queries by Tenant', 2)

    doc.add_paragraph('HealthPlus (Healthcare):')
    healthplus_queries = [
        '"show all patients"',
        '"how many doctors in cardiology"',
        '"appointments this month"',
        '"average treatment cost"',
        '"top 5 patients by number of visits"'
    ]
    for query in healthplus_queries:
        doc.add_paragraph(query, style='List Bullet')

    doc.add_paragraph('FinanceHub (Finance):')
    finance_queries = [
        '"list all accounts"',
        '"total transactions this year"',
        '"accounts with balance over $10000"',
        '"average loan amount"',
        '"top investments by return"'
    ]
    for query in finance_queries:
        doc.add_paragraph(query, style='List Bullet')

    doc.add_paragraph('RetailMax (Retail):')
    retail_queries = [
        '"show all products"',
        '"products under $50"',
        '"best selling products"',
        '"inventory below 10 units"',
        '"total sales by category"'
    ]
    for query in retail_queries:
        doc.add_paragraph(query, style='List Bullet')

    add_heading_with_style(doc, 'Exporting Results', 2)

    export_steps = [
        '1. Execute a query to generate results',
        '2. Click the export button in the UI',
        '3. Select format: CSV, JSON, or Excel',
        '4. File is saved to /exports directory',
        '5. Download link appears in the interface'
    ]
    for step in export_steps:
        doc.add_paragraph(step, style='List Number')

    add_heading_with_style(doc, 'Common Issues and Solutions', 2)

    troubleshooting = [
        ['Issue', 'Solution'],
        ['Cannot login', 'Verify email domain matches organization (e.g., @healthplus.com)'],
        ['Database connection error', 'Check Docker containers are running: docker-compose ps'],
        ['Query returns no results', 'Ensure query syntax matches expected patterns'],
        ['Permission denied', 'Check user role has sufficient permissions for query type'],
        ['Slow query performance', 'Reduce query complexity or contact admin'],
        ['Token expired', 'Re-login to get new JWT token']
    ]
    add_table_from_data(doc, troubleshooting[0], troubleshooting[1:])

    doc.add_page_break()

    # ============================================================================
    # 18. APPENDICES
    # ============================================================================
    add_heading_with_style(doc, '18. Appendices', 1)

    add_heading_with_style(doc, 'Appendix A: File Manifest', 2)
    doc.add_paragraph('Complete listing of key project files:')

    file_categories = [
        'Core Application (35 files in /src)',
        'Tests (9 files in /tests)',
        'Database Schemas (4 tenant init files, RBAC schema)',
        'Root Templates (4 database types with versions)',
        'Configuration (docker-compose.yml, .env, requirements.txt)',
        'Frontend (streamlit_app.py, 3 alternative frontends)',
        'Demo Scripts (4 launch/demo files)',
        'Documentation (This document)'
    ]
    for category in file_categories:
        doc.add_paragraph(category, style='List Bullet')

    add_heading_with_style(doc, 'Appendix B: Database Ports', 2)

    port_reference = [
        ['Service', 'Host Port', 'Container Port', 'Protocol'],
        ['HealthPlus MySQL', '3307', '3306', 'TCP'],
        ['EduLearn MySQL', '3308', '3306', 'TCP'],
        ['FinanceHub PostgreSQL', '5433', '5432', 'TCP'],
        ['RetailMax MongoDB', '27018', '27017', 'TCP'],
        ['FastAPI Backend', '8000', '8000', 'HTTP'],
        ['Streamlit Frontend', '8501', '8501', 'HTTP'],
        ['phpMyAdmin', '8080', '80', 'HTTP'],
        ['pgAdmin', '8081', '80', 'HTTP'],
        ['Mongo Express', '8082', '8081', 'HTTP'],
        ['Ollama LLM', '11434', '11434', 'HTTP']
    ]
    add_table_from_data(doc, port_reference[0], port_reference[1:])

    add_heading_with_style(doc, 'Appendix C: Environment Variable Reference', 2)
    doc.add_paragraph('Complete list of configurable environment variables:')

    env_vars = [
        ['Variable', 'Default', 'Description'],
        ['SECRET_KEY', 'your-super-secret-key-...', 'JWT signing key (MUST change in production)'],
        ['ALGORITHM', 'HS256', 'JWT algorithm'],
        ['ACCESS_TOKEN_EXPIRE_MINUTES', '30', 'Token expiration time'],
        ['DEBUG', 'True', 'Debug mode (set False in production)'],
        ['LOG_LEVEL', 'INFO', 'Logging verbosity'],
        ['MAX_QUERY_RESULTS', '1000', 'Maximum rows returned per query'],
        ['RATE_LIMIT_REQUESTS', '30', 'Max requests per window'],
        ['RATE_LIMIT_WINDOW', '60', 'Rate limit window in seconds'],
        ['OLLAMA_BASE_URL', 'http://localhost:11434', 'Ollama service URL'],
        ['OLLAMA_MODEL', 'llama3.1', 'LLM model name']
    ]
    add_table_from_data(doc, env_vars[0], env_vars[1:])

    add_heading_with_style(doc, 'Appendix D: API Response Codes', 2)

    response_codes = [
        ['Code', 'Status', 'Meaning'],
        ['200', 'OK', 'Request successful'],
        ['201', 'Created', 'Resource created successfully'],
        ['400', 'Bad Request', 'Invalid request parameters'],
        ['401', 'Unauthorized', 'Invalid or missing JWT token'],
        ['403', 'Forbidden', 'Insufficient permissions'],
        ['404', 'Not Found', 'Resource does not exist'],
        ['422', 'Unprocessable Entity', 'Validation error'],
        ['429', 'Too Many Requests', 'Rate limit exceeded'],
        ['500', 'Internal Server Error', 'Server-side error'],
        ['503', 'Service Unavailable', 'Database or service down']
    ]
    add_table_from_data(doc, response_codes[0], response_codes[1:])

    add_heading_with_style(doc, 'Appendix E: Glossary', 2)

    glossary = [
        ['Term', 'Definition'],
        ['HDT', 'Human Digital Twin - Personalized AI assistant profile'],
        ['RBAC', 'Role-Based Access Control - Permission management system'],
        ['JWT', 'JSON Web Token - Authentication token format'],
        ['NLP2SQL', 'Natural Language Processing to SQL conversion'],
        ['Tenant', 'An isolated organization/customer in the multi-tenant system'],
        ['org_id', 'Organization identifier used for data isolation'],
        ['ORM', 'Object-Relational Mapping - SQLAlchemy abstraction layer'],
        ['ASGI', 'Asynchronous Server Gateway Interface - Python web server protocol'],
        ['Ollama', 'Local LLM inference engine'],
        ['Connection Pool', 'Reusable database connections for performance']
    ]
    add_table_from_data(doc, glossary[0], glossary[1:])

    add_heading_with_style(doc, 'Appendix F: Future Enhancements', 2)
    doc.add_paragraph('Potential improvements and roadmap items:')

    future_items = [
        'Advanced NLP with transformer models for better query understanding',
        'Natural language result explanations',
        'Query optimization suggestions',
        'Automated dashboard creation from queries',
        'Real-time query collaboration',
        'Scheduled query execution',
        'Email/Slack notifications for query results',
        'Advanced data visualization with custom charts',
        'Machine learning model integration',
        'GraphQL API in addition to REST',
        'Mobile application (iOS/Android)',
        'Multi-language support (internationalization)',
        'Advanced caching with Redis',
        'Kubernetes deployment configuration',
        'Single Sign-On (SSO) integration'
    ]
    for item in future_items:
        doc.add_paragraph(item, style='List Bullet')

    # ============================================================================
    # DOCUMENT FOOTER
    # ============================================================================
    doc.add_page_break()

    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run(f'\n\n\nEnd of Documentation\n\n')
    footer_para.add_run(f'Multi-Tenant NLP2SQL System v1.0\n')
    footer_para.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}\n\n')
    footer_para.add_run('For questions or support, please contact the development team.')

    # Save document
    output_path = 'D:\\Fundae\\Multi Tenant NLP2SQL\\Multi_Tenant_NLP2SQL_Comprehensive_Documentation.docx'
    doc.save(output_path)
    print(f'\n[SUCCESS] Documentation successfully created at:\n{output_path}\n')
    print(f'[INFO] Total sections: 18')
    print(f'[INFO] Multiple tables and detailed explanations included')
    print(f'[INFO] Estimated pages: 40-50')

    return output_path

if __name__ == '__main__':
    create_documentation()

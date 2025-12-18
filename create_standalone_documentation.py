"""
Script to create comprehensive Word documentation for Multi-Tenant NLP2SQL Standalone System
Specifically for: python -m streamlit run streamlit_standalone.py --server.port 8504
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def add_heading_with_style(doc, text, level=1):
    """Add a styled heading"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_table_from_data(doc, headers, data):
    """Create a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True

    # Data rows
    for row_data in data:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = str(value)

    return table

def create_documentation():
    """Create the comprehensive standalone system documentation"""
    doc = Document()

    # ============================================================================
    # TITLE PAGE
    # ============================================================================
    title = doc.add_heading('Multi-Tenant NLP2SQL', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Standalone Application Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.color.rgb = RGBColor(0, 0, 139)

    doc.add_paragraph()

    launch_cmd = doc.add_paragraph()
    launch_cmd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_text = launch_cmd.add_run('python -m streamlit run streamlit_standalone.py --server.port 8504')
    run_text.font.size = Pt(14)
    run_text.font.name = 'Courier New'
    run_text.font.color.rgb = RGBColor(139, 0, 0)
    run_text.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}\n')
    info_para.add_run('Version: 1.0 (Standalone)\n')
    info_para.add_run('Location: D:\\Fundae\\Multi Tenant NLP2SQL')

    doc.add_page_break()

    # ============================================================================
    # TABLE OF CONTENTS
    # ============================================================================
    add_heading_with_style(doc, 'Table of Contents', 1)

    toc_items = [
        '1. Executive Summary',
        '2. System Overview',
        '3. Standalone Architecture',
        '4. Key Features',
        '5. Quick Start Guide',
        '6. User Roles & Permissions',
        '7. Multi-Tenancy in Standalone Mode',
        '8. Query Processing',
        '9. Database Integration',
        '10. User Interface Components',
        '11. Security & Access Control',
        '12. Demo Accounts',
        '13. Sample Queries by Tenant',
        '14. Data Visualization',
        '15. Troubleshooting',
        '16. Technical Details',
        '17. Configuration',
        '18. Appendices'
    ]

    for item in toc_items:
        doc.add_paragraph(item, style='List Number')

    doc.add_page_break()

    # ============================================================================
    # 1. EXECUTIVE SUMMARY
    # ============================================================================
    add_heading_with_style(doc, '1. Executive Summary', 1)

    doc.add_paragraph(
        'The Multi-Tenant NLP2SQL Standalone Application is a self-contained, lightweight system '
        'that enables users to query databases using natural language - all within a single Streamlit '
        'application. Unlike the full API-based version, this standalone mode requires no backend '
        'server, making it ideal for demos, development, and single-machine deployments.'
    )

    add_heading_with_style(doc, 'What Makes This "Standalone"?', 2)
    standalone_features = [
        'NO Backend API Required - Everything runs in Streamlit',
        'NO FastAPI Server Needed - Direct database access',
        'NO Docker Compose Required - Uses SQLite databases',
        'Single Command Launch - One simple command to run',
        'Built-in Mock Authentication - No external auth service',
        'Embedded Permission System - RBAC built into the app',
        'SQLite-Based - Lightweight database files',
        'Self-Contained - All logic in one Python file (910 lines)'
    ]

    for feature in standalone_features:
        doc.add_paragraph(feature, style='List Bullet')

    add_heading_with_style(doc, 'Quick Stats', 2)
    stats_data = [
        ['Metric', 'Value'],
        ['Application File', 'streamlit_standalone.py (910 lines)'],
        ['Launch Command', 'python -m streamlit run streamlit_standalone.py --server.port 8504'],
        ['Port', '8504'],
        ['Access URL', 'http://localhost:8504'],
        ['Database Type', 'SQLite (file-based)'],
        ['Tenants Supported', '2 (TechCorp, HealthPlus)'],
        ['Demo Users', '3 pre-configured accounts'],
        ['User Roles', '4 (Admin, Analyst, User, Viewer)'],
        ['No External Dependencies', 'Backend, Docker, MySQL, PostgreSQL, MongoDB']
    ]

    add_table_from_data(doc, stats_data[0], stats_data[1:])

    doc.add_page_break()

    # ============================================================================
    # 2. SYSTEM OVERVIEW
    # ============================================================================
    add_heading_with_style(doc, '2. System Overview', 1)

    add_heading_with_style(doc, 'Purpose', 2)
    doc.add_paragraph(
        'This standalone application demonstrates multi-tenant NLP2SQL capabilities in a simplified, '
        'self-contained package. It is perfect for:'
    )

    purposes = [
        'Quick Demonstrations - Show NLP2SQL capabilities without complex setup',
        'Development & Testing - Rapid prototyping and feature testing',
        'Training & Education - Teaching users about the system',
        'Offline Usage - No internet or external services required',
        'Simple Deployments - Single machine, minimal infrastructure'
    ]

    for purpose in purposes:
        doc.add_paragraph(purpose, style='List Bullet')

    add_heading_with_style(doc, 'Architecture Comparison', 2)

    comparison_data = [
        ['Aspect', 'Standalone Mode (Port 8504)', 'Full API Mode (Port 8501)'],
        ['Backend Server', 'None - All in Streamlit', 'FastAPI on port 8000'],
        ['Database', 'SQLite (local files)', 'MySQL, PostgreSQL, MongoDB, SQLite'],
        ['Authentication', 'Mock in-memory', 'JWT with RBAC database'],
        ['Deployment', 'Single Python command', 'Docker Compose multi-container'],
        ['Query Engine', 'Embedded NL to SQL', 'Separate NLP2SQL service'],
        ['File Count', '1 main file', '35+ source files'],
        ['Complexity', 'Low - Easy to understand', 'High - Production-ready'],
        ['Use Case', 'Demo, Dev, Testing', 'Production, Enterprise']
    ]

    add_table_from_data(doc, comparison_data[0], comparison_data[1:])

    doc.add_page_break()

    # ============================================================================
    # 3. STANDALONE ARCHITECTURE
    # ============================================================================
    add_heading_with_style(doc, '3. Standalone Architecture', 1)

    add_heading_with_style(doc, 'Single-File Architecture', 2)
    doc.add_paragraph(
        'The entire standalone system is contained in streamlit_standalone.py (910 lines). '
        'All components are embedded within this single file:'
    )

    components = [
        ['Component', 'Lines', 'Description'],
        ['Imports & Config', '1-62', 'Streamlit setup, role permissions configuration'],
        ['Mock Users', '63-105', 'Pre-configured demo accounts'],
        ['Tenant Data', '106-271', 'Mock data for TechCorp and HealthPlus'],
        ['Mock Queries', '272-364', 'Fallback query responses'],
        ['Authentication', '365-372', 'Simple email/password auth'],
        ['Permission System', '373-459', 'Role-based query access control'],
        ['Database Functions', '460-495', 'SQLite connection and query execution'],
        ['NLP to SQL Engine', '496-569', 'Natural language query parser'],
        ['Query Processor', '570-613', 'Main query execution logic'],
        ['Query Suggestions', '614-634', 'Context-aware query hints'],
        ['Login Page UI', '635-676', 'Authentication interface'],
        ['Main App UI', '677-883', 'Query interface and results display'],
        ['Entry Point', '884-910', 'Main function and CSS styling']
    ]

    add_table_from_data(doc, components[0], components[1:])

    add_heading_with_style(doc, 'Data Flow', 2)
    doc.add_paragraph('Request processing in standalone mode:')

    flow_steps = [
        '1. User logs in via Streamlit login page',
        '2. Credentials checked against MOCK_USERS dictionary',
        '3. User profile stored in st.session_state',
        '4. User enters natural language query',
        '5. Permission check against ROLE_PERMISSIONS',
        '6. Query converted to SQL by generate_sql_from_nl()',
        '7. SQL executed on SQLite database via execute_database_query()',
        '8. Results displayed in Streamlit dataframe',
        '9. Automatic visualization if applicable',
        '10. Export options (CSV, JSON) provided'
    ]

    for step in flow_steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_page_break()

    # ============================================================================
    # 4. KEY FEATURES
    # ============================================================================
    add_heading_with_style(doc, '4. Key Features', 1)

    add_heading_with_style(doc, 'Natural Language Processing', 2)
    doc.add_paragraph(
        'The standalone version includes a pattern-matching NLP engine that converts '
        'natural language queries to SQL:'
    )

    nlp_features = [
        'Product Queries - "show me all products", "products under $100"',
        'User Queries - "how many users", "list all active users"',
        'Customer Analytics - "top customers by sales"',
        'Order Tracking - "recent orders", "pending orders"',
        'Aggregations - "average price by category"',
        'Inventory - "stock levels", "low inventory items"',
        'System Metrics - "system performance" (admin only)',
        'Error Logs - "error log" (admin only)'
    ]

    for feature in nlp_features:
        doc.add_paragraph(feature, style='List Bullet')

    add_heading_with_style(doc, 'Role-Based Access Control', 2)
    doc.add_paragraph(
        'Built-in RBAC system without external database:'
    )

    role_features = [
        'Admin - Full access to all queries including system metrics and user management',
        'Analyst - Data analysis, customer queries, exports (no system access)',
        'User - Basic product/service queries only',
        'Viewer - Read-only product viewing (no exports, no analytics)'
    ]

    for feature in role_features:
        doc.add_paragraph(feature, style='List Bullet')

    add_heading_with_style(doc, 'Multi-Tenancy', 2)
    doc.add_paragraph('Tenant isolation features:')

    mt_features = [
        'Tenant-Specific Data - TechCorp sees tech products, HealthPlus sees medical services',
        'Separate Databases - techcorp_db.sqlite and healthplus_db.sqlite',
        'User-Tenant Mapping - Users assigned to specific organizations',
        'Isolated Query Results - No cross-tenant data leakage',
        'Tenant-Aware UI - Organization name displayed in header'
    ]

    for feature in mt_features:
        doc.add_paragraph(feature, style='List Bullet')

    add_heading_with_style(doc, 'Data Visualization', 2)
    doc.add_paragraph('Automatic chart generation using Plotly:')

    viz_features = [
        'Revenue/Sales Charts - Bar charts for sales data',
        'Price Distribution - Histograms for pricing analysis',
        'Scatter Plots - Correlation analysis for numeric data',
        'Interactive Charts - Zoom, pan, hover tooltips',
        'Export Visualizations - Download charts as images'
    ]

    for feature in viz_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_page_break()

    # ============================================================================
    # 5. QUICK START GUIDE
    # ============================================================================
    add_heading_with_style(doc, '5. Quick Start Guide', 1)

    add_heading_with_style(doc, 'Prerequisites', 2)
    prereqs = [
        'Python 3.9 or higher installed',
        'pip package manager',
        'Internet connection (for initial pip install only)'
    ]

    for prereq in prereqs:
        doc.add_paragraph(prereq, style='List Bullet')

    add_heading_with_style(doc, 'Installation Steps', 2)

    install_steps = [
        '1. Open Command Prompt or Terminal',
        '2. Navigate to project directory:',
        '   cd "D:\\Fundae\\Multi Tenant NLP2SQL"',
        '3. Install required Python packages:',
        '   pip install streamlit pandas plotly',
        '4. Verify streamlit_standalone.py exists in the directory'
    ]

    for step in install_steps:
        doc.add_paragraph(step, style='List Number')

    add_heading_with_style(doc, 'Launch the Application', 2)

    para = doc.add_paragraph('Run the following command:')
    para.add_run('\n\npython -m streamlit run streamlit_standalone.py --server.port 8504\n\n').bold = True

    doc.add_paragraph(
        'The application will start and automatically open in your default web browser at '
        'http://localhost:8504'
    )

    add_heading_with_style(doc, 'First Time Login', 2)

    login_steps = [
        '1. Application loads with login page',
        '2. Expand "Demo Accounts" to see available credentials',
        '3. Use admin@techcorp.com / admin123 for full access',
        '4. Click "Login" button',
        '5. Main application interface appears',
        '6. Enter a natural language query',
        '7. Click "Execute Query" to see results'
    ]

    for step in login_steps:
        doc.add_paragraph(step, style='List Number')

    add_heading_with_style(doc, 'Stopping the Application', 2)
    doc.add_paragraph('To stop the application:')

    stop_steps = [
        '1. Go to the Command Prompt/Terminal window',
        '2. Press Ctrl+C',
        '3. Confirm shutdown if prompted',
        '4. Browser tab can be closed'
    ]

    for step in stop_steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_page_break()

    # ============================================================================
    # 6. USER ROLES & PERMISSIONS
    # ============================================================================
    add_heading_with_style(doc, '6. User Roles & Permissions', 1)

    add_heading_with_style(doc, 'Role Hierarchy', 2)

    roles_data = [
        ['Role', 'Query Access', 'Features', 'Restrictions'],
        ['Admin', 'All queries', 'User management, System metrics, Error logs, Analytics, Exports, Data modification', 'None'],
        ['Analyst', 'Read + Analytics', 'Analytics, Data exports', 'Cannot access: User management, System metrics, Error logs, Modify data'],
        ['User', 'Read only', 'Basic queries', 'Cannot access: Analytics, Customer data, User management, System data, Exports'],
        ['Viewer', 'Read only', 'View products/services', 'Cannot access: Everything except basic product viewing']
    ]

    add_table_from_data(doc, roles_data[0], roles_data[1:])

    add_heading_with_style(doc, 'Permission Enforcement', 2)
    doc.add_paragraph(
        'The check_query_permission() function enforces role-based restrictions by matching '
        'query text against restricted patterns:'
    )

    permission_categories = [
        ['Restriction Category', 'Blocked Patterns', 'Allowed Roles'],
        ['System Metrics', 'system performance, performance metrics, system health', 'Admin only'],
        ['User Management', 'all users, active users, user list, delete/create/modify user', 'Admin only'],
        ['Error Logs', 'error log, system log, debug log', 'Admin only'],
        ['Data Modification', 'delete, drop, truncate, update, insert, alter', 'Admin only'],
        ['Advanced Analytics', 'machine learning, predictive, forecast, trend analysis', 'Analyst, Admin'],
        ['Customer/Sales Data', 'customer, top, sales, revenue, order', 'Analyst, Admin']
    ]

    add_table_from_data(doc, permission_categories[0], permission_categories[1:])

    add_heading_with_style(doc, 'Permission Denied Messages', 2)
    doc.add_paragraph(
        'When a user attempts a restricted query, the system displays:'
    )

    denial_info = [
        'Clear error message explaining why access was denied',
        'Required role to execute the query',
        'Suggestion on what to do next',
        'List of what the user CAN do with their current role',
        'Expandable section showing full role permissions'
    ]

    for info in denial_info:
        doc.add_paragraph(info, style='List Bullet')

    doc.add_page_break()

    # ============================================================================
    # 7. MULTI-TENANCY IN STANDALONE MODE
    # ============================================================================
    add_heading_with_style(doc, '7. Multi-Tenancy in Standalone Mode', 1)

    add_heading_with_style(doc, 'Tenant Organizations', 2)
    doc.add_paragraph('The standalone system supports two tenant organizations:')

    tenant_data = [
        ['Tenant', 'Organization', 'Industry', 'Database File', 'Users'],
        ['techcorp', 'TechCorp Solutions', 'Technology', 'demo_databases/techcorp_db.sqlite', 'admin@techcorp.com, analyst@techcorp.com'],
        ['healthplus', 'HealthPlus Medical', 'Healthcare', 'demo_databases/healthplus_db.sqlite', 'user@healthplus.com']
    ]

    add_table_from_data(doc, tenant_data[0], tenant_data[1:])

    add_heading_with_style(doc, 'Tenant Data Isolation', 2)
    doc.add_paragraph('How tenant isolation works in standalone mode:')

    isolation_methods = [
        '1. User Login - Email domain maps to organization (admin@techcorp.com -> TechCorp)',
        '2. Tenant Assignment - User profile includes org_id field',
        '3. Session Storage - st.session_state.current_tenant tracks active tenant',
        '4. Database Routing - Queries directed to tenant-specific SQLite file',
        '5. Data Filtering - Results only from the user\'s organization',
        '6. Mock Data Separation - TENANT_DATA dictionary segregates data by org_id'
    ]

    for method in isolation_methods:
        doc.add_paragraph(method, style='List Number')

    add_heading_with_style(doc, 'Tenant-Specific Data Examples', 2)

    doc.add_paragraph('TechCorp (Technology Products):')
    techcorp_data = [
        'Machine Learning Platform - $1,299.99',
        'Enterprise Software License - $999.99',
        'Data Warehouse Service - $899.99',
        'Database Management System - $799.99',
        '...20 total technology products'
    ]
    for item in techcorp_data:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph('HealthPlus (Medical Services):')
    healthplus_data = [
        'Minor Surgery - $850.00',
        'MRI Scan - $450.00',
        'CT Scan - $350.00',
        'Emergency Visit - $275.00',
        '...20 total medical services'
    ]
    for item in healthplus_data:
        doc.add_paragraph(item, style='List Bullet')

    add_heading_with_style(doc, 'Tenant Switching', 2)
    doc.add_paragraph(
        'To switch tenants, user must log out and log back in with credentials '
        'from a different organization. No cross-tenant access is permitted within '
        'a single session.'
    )

    doc.add_page_break()

    # ============================================================================
    # 8. QUERY PROCESSING
    # ============================================================================
    add_heading_with_style(doc, '8. Query Processing', 1)

    add_heading_with_style(doc, 'NLP to SQL Conversion', 2)
    doc.add_paragraph(
        'The generate_sql_from_nl() function uses pattern matching to convert '
        'natural language to SQL:'
    )

    conversion_examples = [
        ['Natural Language', 'Generated SQL'],
        ['"show me all products"', 'SELECT * FROM products ORDER BY price DESC'],
        ['"products under $100"', 'SELECT name, price, category, stock_quantity FROM products WHERE price < 100 ORDER BY price DESC'],
        ['"how many users"', 'SELECT COUNT(*) as total_users FROM users'],
        ['"top customers"', 'SELECT customer_name, total_orders, total_spent FROM customers ORDER BY total_spent DESC LIMIT 10'],
        ['"average price by category"', 'SELECT category, AVG(price) as avg_price, COUNT(*) as count FROM products GROUP BY category ORDER BY avg_price DESC'],
        ['"low inventory"', 'SELECT name, category, stock_quantity, price FROM products ORDER BY stock_quantity ASC LIMIT 20']
    ]

    add_table_from_data(doc, conversion_examples[0], conversion_examples[1:])

    add_heading_with_style(doc, 'Query Processing Pipeline', 2)

    pipeline_steps = [
        '1. Query Input - User enters natural language text',
        '2. Permission Check - check_query_permission() validates access',
        '3. Tenant Identification - Extract org_id from user session',
        '4. SQL Generation - generate_sql_from_nl() creates SQL query',
        '5. Database Selection - Identify correct SQLite file for tenant',
        '6. Query Execution - execute_database_query() runs SQL',
        '7. Result Collection - Fetch rows as list of dictionaries',
        '8. Performance Tracking - Record execution time in milliseconds',
        '9. Result Formatting - Convert to Pandas DataFrame',
        '10. Visualization - Auto-generate charts if applicable'
    ]

    for step in pipeline_steps:
        doc.add_paragraph(step, style='List Number')

    add_heading_with_style(doc, 'Supported Query Patterns', 2)

    patterns = [
        ['Pattern Type', 'Keywords', 'Example'],
        ['Product Listing', 'show, product, service, all', '"show me all products"'],
        ['Price Filtering', 'under, less than, below, over, more than', '"products under $50"'],
        ['Counting', 'how many, count', '"how many users"'],
        ['Top N', 'top, best', '"top 10 customers"'],
        ['Aggregation', 'average, avg, sum, total', '"average price by category"'],
        ['Sorting', 'expensive, cheap, recent, latest', '"most expensive products"'],
        ['Status Filter', 'active, pending, completed', '"active users"'],
        ['Inventory', 'stock, inventory, low', '"low inventory items"'],
        ['System Metrics', 'system, performance, metrics', '"system performance"']
    ]

    add_table_from_data(doc, patterns[0], patterns[1:])

    doc.add_page_break()

    # ============================================================================
    # 9. DATABASE INTEGRATION
    # ============================================================================
    add_heading_with_style(doc, '9. Database Integration', 1)

    add_heading_with_style(doc, 'SQLite Databases', 2)
    doc.add_paragraph(
        'The standalone system uses SQLite file-based databases stored in the '
        'demo_databases directory:'
    )

    db_files = [
        'demo_databases/techcorp_db.sqlite - TechCorp technology products',
        'demo_databases/healthplus_db.sqlite - HealthPlus medical services'
    ]

    for db in db_files:
        doc.add_paragraph(db, style='List Bullet')

    add_heading_with_style(doc, 'Database Schema', 2)
    doc.add_paragraph('Common schema across both tenant databases:')

    schema_tables = [
        ['Table', 'Columns', 'Purpose'],
        ['products', 'product_id, name, price, category, stock_quantity, description, created_at', 'Product/service catalog'],
        ['users', 'user_id, username, full_name, email, role, department, status, last_login', 'User accounts'],
        ['customers', 'customer_id, customer_name, total_orders, total_spent, created_at', 'Customer information'],
        ['orders', 'order_id, customer_id, order_date, order_total, status', 'Order tracking']
    ]

    add_table_from_data(doc, schema_tables[0], schema_tables[1:])

    add_heading_with_style(doc, 'Database Connection', 2)
    doc.add_paragraph('The execute_database_query() function handles database operations:')

    db_operations = [
        '1. Resolve database path from tenant ID',
        '2. Check if database file exists',
        '3. Create SQLite connection',
        '4. Set row_factory to sqlite3.Row for dict conversion',
        '5. Execute SQL query',
        '6. Fetch all results',
        '7. Convert rows to list of dictionaries',
        '8. Close connection',
        '9. Return results or empty list on error'
    ]

    for op in db_operations:
        doc.add_paragraph(op, style='List Number')

    add_heading_with_style(doc, 'Fallback Mechanism', 2)
    doc.add_paragraph(
        'If SQLite databases are not found, the system falls back to TENANT_DATA '
        'and MOCK_QUERIES dictionaries containing pre-populated mock data. This ensures '
        'the application always works, even without database files.'
    )

    doc.add_page_break()

    # ============================================================================
    # 10. USER INTERFACE COMPONENTS
    # ============================================================================
    add_heading_with_style(doc, '10. User Interface Components', 1)

    add_heading_with_style(doc, 'Login Page', 2)
    doc.add_paragraph('Features of the login interface:')

    login_features = [
        'Title: "Multi-Tenant NLP2SQL System"',
        'Centered login form with email and password fields',
        'Expandable "Demo Accounts" section showing available credentials',
        'Login button with emoji icon',
        'Error messages for invalid credentials',
        'Warning for missing fields',
        'Gradient background (purple to violet)'
    ]

    for feature in login_features:
        doc.add_paragraph(feature, style='List Bullet')

    add_heading_with_style(doc, 'Main Application Layout', 2)

    layout_sections = [
        ['Section', 'Location', 'Content'],
        ['Header', 'Top', 'Organization name and app title'],
        ['Sidebar (Left)', 'Left panel', 'User profile, logout, query suggestions'],
        ['Query Input', 'Main area (left)', 'Text area for natural language queries'],
        ['Action Buttons', 'Main area', 'Execute Query, Clear buttons'],
        ['Results Display', 'Main area', 'SQL code, metrics, data table, visualizations'],
        ['Export Options', 'Below results', 'Download CSV, JSON buttons'],
        ['Query History', 'Main area (right)', 'Recent queries with re-run option']
    ]

    add_table_from_data(doc, layout_sections[0], layout_sections[1:])

    add_heading_with_style(doc, 'Sidebar Components', 2)

    sidebar_items = [
        'User Profile Display - Full name, role, organization, tenant ID',
        'Logout Button - Red button to sign out',
        'Divider - Visual separator',
        'Quick Queries Section - Up to 5 query suggestions',
        'Clickable Suggestions - Each suggestion is a button that fills the query input'
    ]

    for item in sidebar_items:
        doc.add_paragraph(item, style='List Bullet')

    add_heading_with_style(doc, 'Results Display', 2)
    doc.add_paragraph('When query completes successfully, the interface shows:')

    results_elements = [
        'Section Header: "Query Results" with checkmark',
        'Expandable SQL Code - Shows generated SQL in syntax-highlighted box',
        'Metrics Row - Three columns showing:',
        '  - Execution Time (ms)',
        '  - Rows Returned (count)',
        '  - Status (success/error)',
        'Data Table - Interactive Streamlit dataframe with sorting',
        'Automatic Visualizations - Charts based on data type',
        'Export Section - Download buttons for CSV and JSON',
        'Success Message - Confirmation with tenant name'
    ]

    for element in results_elements:
        doc.add_paragraph(element, style='List Bullet')

    add_heading_with_style(doc, 'Error Display', 2)
    doc.add_paragraph('For permission denied errors:')

    error_display = [
        'Error Header - "Access Denied" in red',
        'Error Message - Explanation of why access was denied',
        'Query Details - Shows attempted query, user role, required role',
        'Suggestion Box - Helpful guidance on what to do',
        'Role Permissions Expander - Shows what user CAN do with current role'
    ]

    for elem in error_display:
        doc.add_paragraph(elem, style='List Bullet')

    doc.add_page_break()

    # ============================================================================
    # 11. SECURITY & ACCESS CONTROL
    # ============================================================================
    add_heading_with_style(doc, '11. Security & Access Control', 1)

    add_heading_with_style(doc, 'Authentication Mechanism', 2)
    doc.add_paragraph(
        'The standalone system uses simple in-memory authentication via the '
        'authenticate_user() function:'
    )

    auth_flow = [
        '1. User submits email and password',
        '2. Email looked up in MOCK_USERS dictionary',
        '3. Password compared directly (plain text comparison)',
        '4. If match: Return success with user profile',
        '5. If no match: Return error message',
        '6. User profile stored in Streamlit session state',
        '7. Session persists until logout or browser close'
    ]

    for step in auth_flow:
        doc.add_paragraph(step, style='List Number')

    add_heading_with_style(doc, 'Security Considerations', 2)

    security_notes = [
        ['Aspect', 'Implementation', 'Production Readiness'],
        ['Passwords', 'Plain text in code', 'NOT SECURE - For demo only'],
        ['Session', 'Streamlit session state', 'Browser session only'],
        ['Token', 'Mock token string', 'No real JWT validation'],
        ['HTTPS', 'Not configured', 'Should use SSL in production'],
        ['SQL Injection', 'Limited by pattern matching', 'Basic protection only'],
        ['Data Storage', 'Local SQLite files', 'No encryption at rest']
    ]

    add_table_from_data(doc, security_notes[0], security_notes[1:])

    add_heading_with_style(doc, 'IMPORTANT SECURITY WARNING', 2)

    warning_para = doc.add_paragraph()
    warning_run = warning_para.add_run(
        'This standalone application is designed for DEMONSTRATION and DEVELOPMENT purposes only. '
        'It should NOT be used in production environments without significant security enhancements:'
    )
    warning_run.font.color.rgb = RGBColor(139, 0, 0)
    warning_run.bold = True

    security_warnings = [
        'Passwords are stored in plain text in the source code',
        'No password hashing (bcrypt, argon2, etc.)',
        'No JWT token validation',
        'No HTTPS/SSL encryption',
        'No rate limiting',
        'No audit logging',
        'No database encryption',
        'No protection against brute force attacks'
    ]

    for warning in security_warnings:
        doc.add_paragraph(warning, style='List Bullet')

    add_heading_with_style(doc, 'Role-Based Restrictions', 2)
    doc.add_paragraph(
        'While authentication is simple, the permission system is robust. The '
        'check_query_permission() function enforces role-based access control by:'
    )

    rbac_enforcement = [
        'Pattern Matching - Detecting restricted keywords in queries',
        'Role Validation - Checking user role against permission matrix',
        'Early Rejection - Blocking query before SQL generation',
        'Informative Errors - Clear messages explaining denial',
        'Suggestion System - Guiding users to allowed operations'
    ]

    for item in rbac_enforcement:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ============================================================================
    # 12. DEMO ACCOUNTS
    # ============================================================================
    add_heading_with_style(doc, '12. Demo Accounts', 1)

    add_heading_with_style(doc, 'Available User Accounts', 2)

    users_data = [
        ['Email', 'Password', 'Role', 'Organization', 'Full Name', 'Description'],
        ['admin@techcorp.com', 'admin123', 'admin', 'TechCorp Solutions', 'Admin User', 'Full system access, can view all data including system metrics'],
        ['analyst@techcorp.com', 'analyst123', 'analyst', 'TechCorp Solutions', 'Data Analyst', 'Analytics and customer data access, no system metrics'],
        ['user@healthplus.com', 'user123', 'user', 'HealthPlus Medical', 'Regular User', 'Basic product/service queries only, limited access']
    ]

    add_table_from_data(doc, users_data[0], users_data[1:])

    add_heading_with_style(doc, 'What Each Account Can Do', 2)

    doc.add_paragraph('Admin Account (admin@techcorp.com):')
    admin_capabilities = [
        'View all products and services',
        'Access customer and sales data',
        'View user lists and user management data',
        'Check system performance metrics',
        'View error logs and system logs',
        'Export data to CSV and JSON',
        'Execute any query without restrictions'
    ]
    for cap in admin_capabilities:
        doc.add_paragraph(cap, style='List Bullet')

    doc.add_paragraph('Analyst Account (analyst@techcorp.com):')
    analyst_capabilities = [
        'View all products and services',
        'Access customer and sales data',
        'Run analytics queries (aggregations, grouping)',
        'Export data to CSV and JSON',
        'Cannot view system metrics',
        'Cannot access user management data',
        'Cannot view error logs'
    ]
    for cap in analyst_capabilities:
        doc.add_paragraph(cap, style='List Bullet')

    doc.add_paragraph('User Account (user@healthplus.com):')
    user_capabilities = [
        'View products and services',
        'Basic filtering (price ranges)',
        'Cannot access customer data',
        'Cannot run analytics queries',
        'Cannot export data',
        'Cannot view system information',
        'Different tenant (HealthPlus) with medical services data'
    ]
    for cap in user_capabilities:
        doc.add_paragraph(cap, style='List Bullet')

    doc.add_page_break()

    # ============================================================================
    # 13. SAMPLE QUERIES BY TENANT
    # ============================================================================
    add_heading_with_style(doc, '13. Sample Queries by Tenant', 1)

    add_heading_with_style(doc, 'TechCorp Solutions Queries', 2)
    doc.add_paragraph('Recommended queries for TechCorp users:')

    techcorp_queries = [
        ['Query', 'Expected Results', 'Access Level'],
        ['show me all products', '20 technology products sorted by price', 'All roles'],
        ['products under $500', 'Technology products under $500', 'All roles'],
        ['most expensive products', 'Top 10 products by price', 'All roles'],
        ['how many users', 'Total user count', 'All roles (basic count)'],
        ['top customers', 'Top 10 customers by spending', 'Analyst, Admin'],
        ['all users', 'Complete user list with roles', 'Admin only'],
        ['active users', 'Active users with last login', 'Admin only'],
        ['system performance', 'CPU, memory, connection metrics', 'Admin only'],
        ['error log', 'Recent system errors and warnings', 'Admin only']
    ]

    add_table_from_data(doc, techcorp_queries[0], techcorp_queries[1:])

    add_heading_with_style(doc, 'HealthPlus Medical Queries', 2)
    doc.add_paragraph('Recommended queries for HealthPlus users:')

    healthplus_queries = [
        ['Query', 'Expected Results', 'Access Level'],
        ['show me all products', '20 medical services sorted by price', 'All roles'],
        ['services under $200', 'Medical services under $200', 'All roles'],
        ['most expensive services', 'Top 10 services by price', 'All roles'],
        ['how many users', 'Total medical staff count', 'All roles (basic count)'],
        ['top customers', 'Top medical facilities by spending', 'Analyst, Admin (blocked for User)'],
        ['all users', 'Complete medical staff list', 'Admin only (blocked for User)'],
        ['system performance', 'System health metrics', 'Admin only (blocked for User)']
    ]

    add_table_from_data(doc, healthplus_queries[0], healthplus_queries[1:])

    add_heading_with_style(doc, 'Advanced Query Examples', 2)

    advanced_queries = [
        ['Query Type', 'Example Query', 'Required Role'],
        ['Price Range', 'show products between $100 and $500', 'All'],
        ['Category Aggregation', 'average price by category', 'Analyst, Admin'],
        ['Inventory Management', 'products with low stock', 'All'],
        ['Department Analysis', 'users by department', 'Admin'],
        ['Order Statistics', 'recent completed orders', 'Analyst, Admin'],
        ['Customer Segmentation', 'customers with over 20 orders', 'Analyst, Admin']
    ]

    add_table_from_data(doc, advanced_queries[0], advanced_queries[1:])

    doc.add_page_break()

    # ============================================================================
    # 14. DATA VISUALIZATION
    # ============================================================================
    add_heading_with_style(doc, '14. Data Visualization', 1)

    add_heading_with_style(doc, 'Automatic Chart Generation', 2)
    doc.add_paragraph(
        'The standalone application automatically generates visualizations based on '
        'the structure of query results:'
    )

    viz_rules = [
        ['Data Pattern', 'Chart Type', 'Condition'],
        ['Revenue/Sales columns', 'Bar Chart', 'Column named "revenue" or "total_sales" exists'],
        ['Price data', 'Histogram', 'Column named "price" exists'],
        ['2+ numeric columns', 'Scatter Plot', 'At least 2 numeric columns in results'],
        ['Time series', 'Line Chart', 'Date/time column with numeric values'],
        ['Category counts', 'Bar Chart', 'Category column with COUNT aggregation']
    ]

    add_table_from_data(doc, viz_rules[0], viz_rules[1:])

    add_heading_with_style(doc, 'Plotly Features', 2)
    doc.add_paragraph('All visualizations use Plotly for interactivity:')

    plotly_features = [
        'Hover Tooltips - Show exact values on mouse hover',
        'Zoom - Click and drag to zoom into chart areas',
        'Pan - Move around the chart while zoomed',
        'Toggle Series - Click legend items to show/hide',
        'Download - Built-in download as PNG button',
        'Responsive - Charts resize with browser window',
        'Professional Styling - Clean, modern appearance'
    ]

    for feature in plotly_features:
        doc.add_paragraph(feature, style='List Bullet')

    add_heading_with_style(doc, 'Visualization Examples', 2)

    doc.add_paragraph('Query: "top customers"')
    doc.add_paragraph('Generates: Bar chart of customers vs. total spent', style='List Bullet')

    doc.add_paragraph('Query: "show me all products"')
    doc.add_paragraph('Generates: Histogram of price distribution across products', style='List Bullet')

    doc.add_paragraph('Query: "average price by category"')
    doc.add_paragraph('Generates: Bar chart of categories vs. average price', style='List Bullet')

    add_heading_with_style(doc, 'Export Options', 2)
    doc.add_paragraph('Results can be exported in multiple formats:')

    export_formats = [
        ['Format', 'File Type', 'Use Case'],
        ['CSV', '.csv', 'Excel, data analysis tools, spreadsheets'],
        ['JSON', '.json', 'Programming, APIs, web applications'],
        ['Chart PNG', '.png (via Plotly)', 'Reports, presentations, documentation']
    ]

    add_table_from_data(doc, export_formats[0], export_formats[1:])

    doc.add_page_break()

    # ============================================================================
    # 15. TROUBLESHOOTING
    # ============================================================================
    add_heading_with_style(doc, '15. Troubleshooting', 1)

    add_heading_with_style(doc, 'Common Issues and Solutions', 2)

    troubleshooting_table = [
        ['Issue', 'Possible Cause', 'Solution'],
        ['Cannot login', 'Wrong credentials', 'Use exactly: admin@techcorp.com / admin123 (case-sensitive)'],
        ['Page won\'t load', 'Wrong URL', 'Ensure using http://localhost:8504 (note port 8504)'],
        ['"Module not found" error', 'Missing packages', 'Run: pip install streamlit pandas plotly'],
        ['Database not found', 'SQLite files missing', 'Run demo_simple.py first or use mock data fallback'],
        ['Query returns nothing', 'Database empty or no match', 'Check if demo data is populated'],
        ['Permission denied', 'Insufficient role', 'Login as admin@techcorp.com for full access'],
        ['Port already in use', 'Another app on 8504', 'Change port: --server.port 8505'],
        ['Slow performance', 'Large dataset', 'Use LIMIT in queries or filter by category'],
        ['Charts not showing', 'No numeric data', 'Query must return numeric columns for visualization'],
        ['Cannot stop app', 'Ctrl+C not working', 'Close terminal window and kill python process']
    ]

    add_table_from_data(doc, troubleshooting_table[0], troubleshooting_table[1:])

    add_heading_with_style(doc, 'Application Won\'t Start', 2)

    startup_checks = [
        '1. Verify Python version: python --version (should be 3.9+)',
        '2. Check if Streamlit is installed: streamlit --version',
        '3. Ensure file exists: dir streamlit_standalone.py',
        '4. Check for syntax errors: python -m py_compile streamlit_standalone.py',
        '5. Try alternate launch: streamlit run streamlit_standalone.py',
        '6. Check firewall: Allow Python through Windows Firewall'
    ]

    for check in startup_checks:
        doc.add_paragraph(check, style='List Number')

    add_heading_with_style(doc, 'Query Not Recognized', 2)
    doc.add_paragraph('If your natural language query isn\'t working:')

    query_tips = [
        'Use simple, direct language: "show me all products"',
        'Include key words: "products", "users", "customers", "orders"',
        'For filtering: "under $100", "over $500", "between"',
        'For aggregation: "average", "total", "count"',
        'For sorting: "top", "expensive", "recent", "latest"',
        'Check role permissions - some queries need Admin role',
        'Try a suggested query from the sidebar first'
    ]

    for tip in query_tips:
        doc.add_paragraph(tip, style='List Bullet')

    add_heading_with_style(doc, 'Getting Help', 2)

    help_resources = [
        'Check the sidebar "Quick Queries" for working examples',
        'Try the demo accounts section on login page',
        'Review the expandable "Generated SQL" to see what was executed',
        'Look at role permissions in the "What can I do?" expander',
        'Check Command Prompt/Terminal for Python error messages'
    ]

    for resource in help_resources:
        doc.add_paragraph(resource, style='List Bullet')

    doc.add_page_break()

    # ============================================================================
    # 16. TECHNICAL DETAILS
    # ============================================================================
    add_heading_with_style(doc, '16. Technical Details', 1)

    add_heading_with_style(doc, 'Technology Stack', 2)

    tech_stack = [
        ['Component', 'Technology', 'Version', 'Purpose'],
        ['Framework', 'Streamlit', '1.28+', 'Web application framework'],
        ['Data Processing', 'Pandas', '2.0+', 'DataFrame operations'],
        ['Visualization', 'Plotly', '5.0+', 'Interactive charts'],
        ['Database', 'SQLite3', 'Built-in', 'Local database'],
        ['Language', 'Python', '3.9+', 'Core programming language']
    ]

    add_table_from_data(doc, tech_stack[0], tech_stack[1:])

    add_heading_with_style(doc, 'File Structure', 2)

    file_structure = [
        ['File/Directory', 'Purpose'],
        ['streamlit_standalone.py', 'Main application file (910 lines)'],
        ['demo_databases/', 'Directory for SQLite database files'],
        ['demo_databases/techcorp_db.sqlite', 'TechCorp tenant database'],
        ['demo_databases/healthplus_db.sqlite', 'HealthPlus tenant database'],
        ['requirements.txt', 'Python package dependencies']
    ]

    add_table_from_data(doc, file_structure[0], file_structure[1:])

    add_heading_with_style(doc, 'Key Functions', 2)

    functions = [
        ['Function', 'Lines', 'Purpose'],
        ['authenticate_user()', '366-372', 'Validate email/password against MOCK_USERS'],
        ['check_query_permission()', '374-459', 'Enforce role-based access control'],
        ['get_database_path()', '461-473', 'Resolve tenant database file path'],
        ['execute_database_query()', '475-495', 'Execute SQL on SQLite database'],
        ['generate_sql_from_nl()', '497-569', 'Convert natural language to SQL'],
        ['process_query()', '571-613', 'Main query processing orchestration'],
        ['get_query_suggestions()', '615-634', 'Generate context-aware query hints'],
        ['login_page()', '636-676', 'Render login interface'],
        ['main_app()', '678-883', 'Render main application interface'],
        ['main()', '885-910', 'Application entry point and routing']
    ]

    add_table_from_data(doc, functions[0], functions[1:])

    add_heading_with_style(doc, 'Session State Variables', 2)

    session_vars = [
        ['Variable', 'Type', 'Purpose'],
        ['authenticated', 'Boolean', 'Whether user is logged in'],
        ['user_info', 'Dict', 'Current user profile and role'],
        ['chat_history', 'List', 'Query history for current session'],
        ['current_tenant', 'String', 'Active tenant organization ID'],
        ['last_result', 'Dict', 'Most recent query result'],
        ['current_query', 'String', 'Query text in input field']
    ]

    add_table_from_data(doc, session_vars[0], session_vars[1:])

    add_heading_with_style(doc, 'Performance Characteristics', 2)

    performance = [
        ['Metric', 'Typical Value', 'Notes'],
        ['Startup Time', '2-5 seconds', 'Initial Streamlit load'],
        ['Query Execution', '10-100 ms', 'SQLite queries on small datasets'],
        ['Page Load', '<1 second', 'After initial startup'],
        ['Memory Usage', '50-150 MB', 'Depends on result set size'],
        ['Concurrent Users', '1 (local)', 'Not designed for multi-user']
    ]

    add_table_from_data(doc, performance[0], performance[1:])

    doc.add_page_break()

    # ============================================================================
    # 17. CONFIGURATION
    # ============================================================================
    add_heading_with_style(doc, '17. Configuration', 1)

    add_heading_with_style(doc, 'Streamlit Configuration', 2)
    doc.add_paragraph('Key Streamlit settings in st.set_page_config():')

    streamlit_config = [
        ['Setting', 'Value', 'Purpose'],
        ['page_title', '"Multi-Tenant NLP2SQL Demo"', 'Browser tab title'],
        ['page_icon', '🤖', 'Browser tab icon'],
        ['layout', '"wide"', 'Use full browser width'],
        ['initial_sidebar_state', '"expanded"', 'Show sidebar on load']
    ]

    add_table_from_data(doc, streamlit_config[0], streamlit_config[1:])

    add_heading_with_style(doc, 'Customizing Port', 2)
    doc.add_paragraph('To run on a different port:')

    port_examples = [
        'Port 8504 (default): python -m streamlit run streamlit_standalone.py --server.port 8504',
        'Port 8080: python -m streamlit run streamlit_standalone.py --server.port 8080',
        'Port 9000: python -m streamlit run streamlit_standalone.py --server.port 9000'
    ]

    for example in port_examples:
        doc.add_paragraph(example, style='List Bullet')

    add_heading_with_style(doc, 'Adding New Users', 2)
    doc.add_paragraph('To add demo users, edit MOCK_USERS dictionary in streamlit_standalone.py:')

    code_para = doc.add_paragraph()
    code_para.add_run(
        'MOCK_USERS = {\n'
        '    "newuser@company.com": {\n'
        '        "password": "password123",\n'
        '        "user_id": "user_004",\n'
        '        "role": "analyst",\n'
        '        "org_id": "techcorp",\n'
        '        "org_name": "TechCorp Solutions",\n'
        '        "full_name": "New User"\n'
        '    }\n'
        '}'
    ).font.name = 'Courier New'

    add_heading_with_style(doc, 'Adding New Tenants', 2)
    doc.add_paragraph('To add a new tenant organization:')

    tenant_steps = [
        '1. Add organization to MOCK_USERS with unique org_id',
        '2. Create new SQLite database: demo_databases/newtenant_db.sqlite',
        '3. Add database path mapping in get_database_path() function',
        '4. Add tenant-specific data to TENANT_DATA dictionary',
        '5. Create users with the new org_id'
    ]

    for step in tenant_steps:
        doc.add_paragraph(step, style='List Number')

    add_heading_with_style(doc, 'Modifying Query Patterns', 2)
    doc.add_paragraph('To customize NLP query patterns, edit generate_sql_from_nl() function:')

    pattern_steps = [
        '1. Identify the query type (product, user, customer, etc.)',
        '2. Add new elif condition checking for keywords',
        '3. Return appropriate SQL SELECT statement',
        '4. Test with various natural language phrasings',
        '5. Add to query suggestions if commonly used'
    ]

    for step in pattern_steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_page_break()

    # ============================================================================
    # 18. APPENDICES
    # ============================================================================
    add_heading_with_style(doc, '18. Appendices', 1)

    add_heading_with_style(doc, 'Appendix A: Complete Launch Commands', 2)

    commands_table = [
        ['Purpose', 'Command'],
        ['Standard Launch', 'python -m streamlit run streamlit_standalone.py --server.port 8504'],
        ['Different Port', 'python -m streamlit run streamlit_standalone.py --server.port 8080'],
        ['Auto-open Browser', 'streamlit run streamlit_standalone.py --server.port 8504 --server.headless false'],
        ['Disable Auto-open', 'streamlit run streamlit_standalone.py --server.port 8504 --server.headless true'],
        ['Custom Address', 'streamlit run streamlit_standalone.py --server.address 0.0.0.0 --server.port 8504']
    ]

    add_table_from_data(doc, commands_table[0], commands_table[1:])

    add_heading_with_style(doc, 'Appendix B: All Demo Credentials', 2)

    all_creds = [
        ['Email', 'Password', 'Organization', 'Role'],
        ['admin@techcorp.com', 'admin123', 'TechCorp Solutions', 'Admin'],
        ['analyst@techcorp.com', 'analyst123', 'TechCorp Solutions', 'Analyst'],
        ['user@healthplus.com', 'user123', 'HealthPlus Medical', 'User']
    ]

    add_table_from_data(doc, all_creds[0], all_creds[1:])

    add_heading_with_style(doc, 'Appendix C: Keyboard Shortcuts', 2)

    shortcuts = [
        ['Shortcut', 'Action'],
        ['Ctrl + C', 'Stop the application (in terminal)'],
        ['Ctrl + R', 'Reload page in browser'],
        ['Ctrl + Shift + R', 'Hard reload (clear cache)'],
        ['F11', 'Toggle fullscreen'],
        ['Ctrl + F', 'Find on page']
    ]

    add_table_from_data(doc, shortcuts[0], shortcuts[1:])

    add_heading_with_style(doc, 'Appendix D: Comparison with Full System', 2)

    comparison_full = [
        ['Feature', 'Standalone (Port 8504)', 'Full System (Port 8501)'],
        ['Backend API', 'None', 'FastAPI on port 8000'],
        ['Authentication', 'Mock in-memory', 'JWT with MySQL RBAC database'],
        ['Databases', 'SQLite only', 'MySQL, PostgreSQL, MongoDB, SQLite'],
        ['Tenants', '2 (TechCorp, HealthPlus)', '5+ (includes FinanceHub, RetailMax, EduLearn)'],
        ['NLP Engine', 'Simple pattern matching', 'Advanced pattern matching + Ollama LLM'],
        ['Connection Pooling', 'No', 'Yes - per tenant'],
        ['Caching', 'No', 'Multi-level caching'],
        ['Monitoring', 'No', 'Full monitoring and alerting'],
        ['Audit Logging', 'No', 'Complete audit trail'],
        ['Docker Required', 'No', 'Yes - docker-compose'],
        ['Setup Complexity', 'Very Low', 'High'],
        ['Production Ready', 'No - Demo only', 'Yes'],
        ['Lines of Code', '910 (single file)', '10,000+ (35+ files)']
    ]

    add_table_from_data(doc, comparison_full[0], comparison_full[1:])

    add_heading_with_style(doc, 'Appendix E: Error Messages Reference', 2)

    error_messages = [
        ['Error Message', 'Cause', 'Solution'],
        ['"Invalid credentials"', 'Wrong email or password', 'Check credentials exactly as shown in demo accounts'],
        ['"Database not found"', 'SQLite file missing', 'Run demo_simple.py or use mock data'],
        ['"Access Denied"', 'Role lacks permission', 'Login as Admin or use allowed query'],
        ['"Database error"', 'SQL execution failed', 'Check generated SQL in expander'],
        ['"Module not found: streamlit"', 'Streamlit not installed', 'pip install streamlit'],
        ['"Port already in use"', 'Port 8504 busy', 'Use different port or kill existing process']
    ]

    add_table_from_data(doc, error_messages[0], error_messages[1:])

    add_heading_with_style(doc, 'Appendix F: Quick Reference Card', 2)

    doc.add_paragraph('Launch Application:')
    launch = doc.add_paragraph()
    launch.add_run('python -m streamlit run streamlit_standalone.py --server.port 8504').bold = True

    doc.add_paragraph('\nAccess URL:')
    url = doc.add_paragraph()
    url.add_run('http://localhost:8504').bold = True

    doc.add_paragraph('\nAdmin Login:')
    admin = doc.add_paragraph()
    admin.add_run('admin@techcorp.com / admin123').bold = True

    doc.add_paragraph('\nSample Query:')
    query = doc.add_paragraph()
    query.add_run('show me all products').bold = True

    doc.add_paragraph('\nStop Application:')
    stop = doc.add_paragraph()
    stop.add_run('Press Ctrl+C in terminal').bold = True

    # ============================================================================
    # DOCUMENT FOOTER
    # ============================================================================
    doc.add_page_break()

    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run(f'\n\n\nEnd of Standalone System Documentation\n\n')
    footer_para.add_run(f'Multi-Tenant NLP2SQL Standalone Application\n')
    footer_para.add_run(f'Running on: python -m streamlit run streamlit_standalone.py --server.port 8504\n\n')
    footer_para.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}\n\n')
    footer_para.add_run('For questions or support, please contact the development team.')

    # Save document
    output_path = 'D:\\Fundae\\Multi Tenant NLP2SQL\\Multi_Tenant_NLP2SQL_Standalone_Documentation.docx'
    doc.save(output_path)
    print(f'\n[SUCCESS] Standalone documentation successfully created at:\n{output_path}\n')
    print(f'[INFO] Total sections: 18')
    print(f'[INFO] Comprehensive tables and detailed explanations included')
    print(f'[INFO] Estimated pages: 35-45')
    print(f'[INFO] Covers: streamlit_standalone.py on port 8504')

    return output_path

if __name__ == '__main__':
    create_documentation()
